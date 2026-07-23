"""Build the polished Word version of the verified EduMatrix manuscript."""

from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
SOURCE = ROOT / "outputs" / "EduMatrix_完整技术文档总稿.md"
OUTPUT = ROOT / "outputs" / "EduMatrix_评委最终技术文档.docx"
SCREENSHOTS = ROOT / "outputs" / "evaluator_screenshots"

FONT = "Noto Sans SC"
HEADING_FONT = "Noto Sans SC"
CODE_FONT = "Consolas"
BODY_SIZE = 11.5
TABLE_SIZE = 9.2
IMAGE_WIDTH = Inches(6.35)
GREEN = "2F5D50"
LIGHT_GREEN = "EAF3EF"
LIGHT_BLUE = "EAF2F8"


def set_run_font(run, *, size: float = BODY_SIZE, bold: bool = False, italic: bool = False, color: str | None = None, code: bool = False, heading: bool = False) -> None:
    family = CODE_FONT if code else HEADING_FONT if heading else FONT
    run.font.name = family
    rfonts = run._element.get_or_add_rPr().get_or_add_rFonts()
    for script in ("ascii", "hAnsi", "eastAsia", "cs"):
        rfonts.set(qn(f"w:{script}"), family)
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def add_inline(paragraph, text: str, *, size: float = BODY_SIZE) -> None:
    """Render the small Markdown subset used by the manuscript."""
    pattern = re.compile(r"(\*\*.+?\*\*|`.+?`|\[.+?\]\(.+?\)|\*.+?\*)")
    cursor = 0
    for match in pattern.finditer(text):
        if match.start() > cursor:
            run = paragraph.add_run(text[cursor:match.start()])
            set_run_font(run, size=size)
        token = match.group(0)
        if token.startswith("**") and token.endswith("**"):
            run = paragraph.add_run(token[2:-2])
            set_run_font(run, size=size, bold=True)
        elif token.startswith("`") and token.endswith("`"):
            run = paragraph.add_run(token[1:-1])
            set_run_font(run, size=max(8.5, size - 1), code=True, color="334E45")
        elif token.startswith("["):
            label, url = re.match(r"\[(.+?)\]\((.+?)\)", token).groups()
            run = paragraph.add_run(label)
            set_run_font(run, size=size, color="1769AA")
            run.underline = True
        elif token.startswith("*") and token.endswith("*"):
            run = paragraph.add_run(token[1:-1])
            set_run_font(run, size=size, italic=True)
        else:
            run = paragraph.add_run(token)
            set_run_font(run, size=size)
        cursor = match.end()
    if cursor < len(text):
        run = paragraph.add_run(text[cursor:])
        set_run_font(run, size=size)


def set_cell_text(cell, text: str, *, size: float = TABLE_SIZE) -> None:
    cell.text = ""
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    add_inline(p, text.strip(), size=size)


def shade_cell(cell, fill: str) -> None:
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), fill)
    cell._tc.get_or_add_tcPr().append(shading)


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def add_page_number(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run("第 ")
    set_run_font(run, size=8.5, color="6B7280")
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = "PAGE"
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run._r.extend([fld_char1, instr_text, fld_char2])
    tail = paragraph.add_run(" 页  |  EduMatrix 智教矩阵")
    set_run_font(tail, size=8.5, color="6B7280")


def add_toc(paragraph) -> None:
    run = paragraph.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = 'TOC \\o "1-3" \\h \\z \\u'
    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "打开 Word 后右键更新目录"
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.extend([fld_begin, instr, fld_sep, text, fld_end])


def add_code(doc: Document, code: str) -> None:
    p = doc.add_paragraph(style="No Spacing")
    p.paragraph_format.left_indent = Inches(0.25)
    p.paragraph_format.right_indent = Inches(0.25)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.0
    run = p.add_run(code.rstrip())
    set_run_font(run, size=8.5, code=True, color="27343A")
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), "F1F5F3")
    p._p.get_or_add_pPr().append(shading)


def split_table_row(line: str) -> list[str]:
    value = line.strip().strip("|")
    cells: list[str] = []
    current: list[str] = []
    escaped = False
    for char in value:
        if char == "|" and not escaped:
            cells.append("".join(current).replace("\\|", "|").strip())
            current = []
        else:
            current.append(char)
        escaped = char == "\\" and not escaped
    cells.append("".join(current).replace("\\|", "|").strip())
    return cells


def is_table_separator(line: str) -> bool:
    cells = split_table_row(line)
    return bool(cells) and all(re.fullmatch(r":?-{3,}:?", cell.replace(" ", "")) for cell in cells)


def add_table(doc: Document, headers: list[str], rows: list[list[str]]) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Light Shading Accent 1"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    for col, value in enumerate(headers):
        set_cell_text(table.rows[0].cells[col], value)
        shade_cell(table.rows[0].cells[col], LIGHT_GREEN)
        for paragraph in table.rows[0].cells[col].paragraphs:
            for run in paragraph.runs:
                run.bold = True
    set_repeat_table_header(table.rows[0])
    for row in rows:
        cells = table.add_row().cells
        for col in range(len(headers)):
            set_cell_text(cells[col], row[col] if col < len(row) else "")
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


def parse_markdown(doc: Document, lines: list[str]) -> None:
    index = 0
    in_code = False
    code_lines: list[str] = []
    while index < len(lines):
        line = lines[index].rstrip("\n")
        if line.startswith("```"):
            if in_code:
                add_code(doc, "\n".join(code_lines))
                code_lines = []
                in_code = False
            else:
                in_code = True
            index += 1
            continue
        if in_code:
            code_lines.append(line)
            index += 1
            continue
        if not line.strip() or line.strip() == "---":
            index += 1
            continue
        image_match = re.match(r"!\[(.*?)\]\((.*?)\)", line)
        if image_match:
            caption, relative_path = image_match.groups()
            image_path = SOURCE.parent / relative_path
            if image_path.exists():
                doc.add_picture(str(image_path), width=Inches(6.65))
                caption_paragraph = doc.add_paragraph()
                caption_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                add_inline(caption_paragraph, caption, size=8.8)
            index += 1
            continue
        heading = re.match(r"^(#{1,6})\s+(.*)$", line)
        if heading:
            level = min(len(heading.group(1)), 3)
            doc.add_heading(heading.group(2).strip(), level=level)
            index += 1
            continue
        if line.startswith(">"):
            p = doc.add_paragraph(style="Intense Quote")
            add_inline(p, line.lstrip("> ").strip(), size=10.5)
            index += 1
            continue
        if line.startswith("- ") or line.startswith("* "):
            p = doc.add_paragraph(style="List Bullet")
            add_inline(p, line[2:].strip(), size=11.0)
            index += 1
            continue
        if re.match(r"^\d+\.\s+", line):
            p = doc.add_paragraph(style="List Number")
            add_inline(p, re.sub(r"^\d+\.\s+", "", line), size=11.0)
            index += 1
            continue
        if line.startswith("|") and index + 1 < len(lines) and is_table_separator(lines[index + 1]):
            headers = split_table_row(line)
            rows: list[list[str]] = []
            index += 2
            while index < len(lines) and lines[index].startswith("|"):
                rows.append(split_table_row(lines[index]))
                index += 1
            add_table(doc, headers, rows)
            continue
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(5)
        p.paragraph_format.line_spacing = 1.45
        add_inline(p, line)
        index += 1


def configure_document(doc: Document) -> None:
    section = doc.sections[0]
    # A4 report page with readable margins; this is a formal technical-report
    # layout rather than a government document template with a mandatory font.
    section.page_width = Inches(8.27)
    section.page_height = Inches(11.69)
    section.top_margin = Inches(0.90)
    section.bottom_margin = Inches(0.85)
    section.left_margin = Inches(0.88)
    section.right_margin = Inches(0.88)
    section.header_distance = Inches(0.35)
    section.footer_distance = Inches(0.35)
    styles = doc.styles
    styles["Normal"].font.name = FONT
    normal_rfonts = styles["Normal"]._element.rPr.get_or_add_rFonts()
    for script in ("ascii", "hAnsi", "eastAsia", "cs"):
        normal_rfonts.set(qn(f"w:{script}"), FONT)
    styles["Normal"].font.size = Pt(BODY_SIZE)
    styles["Normal"].paragraph_format.line_spacing = 1.45
    styles["Normal"].paragraph_format.space_after = Pt(5)
    for name, size, color in (
        ("Title", 34, "19352F"),
        ("Heading 1", 20, "304B3D"),
        ("Heading 2", 16, "3F6B55"),
        ("Heading 3", 13.5, "4D6C5B"),
    ):
        style = styles[name]
        style.font.name = HEADING_FONT
        heading_rfonts = style._element.rPr.get_or_add_rFonts()
        for script in ("ascii", "hAnsi", "eastAsia", "cs"):
            heading_rfonts.set(qn(f"w:{script}"), HEADING_FONT)
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.keep_with_next = True
        style.paragraph_format.space_before = Pt(12 if name != "Heading 3" else 8)
        style.paragraph_format.space_after = Pt(6)
    settings = doc.settings._element
    update_fields = settings.find(qn("w:updateFields"))
    if update_fields is None:
        update_fields = OxmlElement("w:updateFields")
        settings.append(update_fields)
    update_fields.set(qn("w:val"), "true")
    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = header.add_run("EduMatrix 智教矩阵  |  最终技术说明版 V2.0")
    set_run_font(run, size=8.2, color="6B7280")
    footer = section.footer.paragraphs[0]
    add_page_number(footer)


def add_cover(doc: Document) -> None:
    title = doc.add_paragraph(style="Title")
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_before = Pt(24)
    title.paragraph_format.space_after = Pt(10)
    title.paragraph_format.line_spacing = 1.05
    title_run = title.add_run("EduMatrix 智教矩阵\n完整技术文档")
    set_run_font(title_run, size=34, bold=True, color="19352F", heading=True)
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_after = Pt(10)
    subtitle.paragraph_format.line_spacing = 1.15
    run = subtitle.add_run("A3：基于大模型的个性化资源生成与学习多智能体系统开发")
    set_run_font(run, size=17, bold=True, color=GREEN, heading=True)
    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.paragraph_format.space_after = Pt(12)
    add_inline(meta, "版本：V2.0  |  验证基线：74f8f271  |  默认核心验收：Windows 原生、无需 Docker", size=11)

    table = doc.add_table(rows=0, cols=2)
    table.style = "Light Shading Accent 1"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    for key, value in (
        ("队伍名称", "上杉绘梨衣"),
        ("参赛学校", "西交利物浦大学"),
        ("队伍成员", "林之正、杨一鸿"),
        ("指导教师", "王成玉"),
        ("文档用途", "作品技术说明、系统验收、评委复现"),
        ("数据口径", "源码事实、运行证据、合成数据明确区分"),
    ):
        cells = table.add_row().cells
        set_cell_text(cells[0], key, size=10.5)
        set_cell_text(cells[1], value, size=10.5)
        cells[0].width = Inches(1.55)
        cells[1].width = Inches(4.85)
        shade_cell(cells[0], LIGHT_BLUE)
        for paragraph in cells[0].paragraphs:
            for r in paragraph.runs:
                r.bold = True
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(2)
    note = doc.add_paragraph(style="Intense Quote")
    note.paragraph_format.line_spacing = 1.25
    add_inline(note, "本文按最终代码和测试证据撰写。视频、技术文档和 PPT 按队伍确认的官方口径不作为本次必交材料；本文件保留为技术说明和评委验收辅助材料。", size=10.5)
    doc.add_page_break()
    doc.add_heading("目录", level=1)
    add_toc(doc.add_paragraph())
    note = doc.add_paragraph(style="Intense Quote")
    add_inline(note, "打开 Word 后右键目录并选择“更新域”，即可刷新页码。", size=9.5)
    doc.add_page_break()


def build() -> None:
    doc = Document()
    configure_document(doc)
    doc.core_properties.title = "EduMatrix 智教矩阵完整技术文档"
    doc.core_properties.subject = "A3 基于大模型的个性化资源生成与学习多智能体系统"
    doc.core_properties.author = "上杉绘梨衣"
    doc.core_properties.keywords = "EduMatrix, 多智能体, 个性化学习, RAG, 软件杯"
    add_cover(doc)

    lines = SOURCE.read_text(encoding="utf-8").splitlines()
    # The cover already contains the manuscript title and document-control page.
    start = next((i for i, line in enumerate(lines) if line.startswith("## 1. 项目摘要")), 0)
    parse_markdown(doc, lines[start:])

    doc.add_page_break()
    doc.add_heading("附录：浏览器验收截图", level=1)
    for filename, caption in (
        ("01-login.png", "图 B-1：登录与注册页面"),
        ("02-dashboard.png", "图 B-2：学习仪表盘"),
        ("03-chat-empty.png", "图 B-3：智能对话入口"),
        ("04-chat-response.png", "图 B-4：多智能体对话与资源结果"),
        ("05-code-sandbox-disabled.png", "图 B-5：代码执行模式边界提示"),
        ("06-learning-path.png", "图 B-6：学习路径可视化"),
        ("07-knowledge.png", "图 B-7：课程知识库"),
        ("08-student-analysis.png", "图 B-8：学习画像与分析"),
        ("09-settings-multimodal.png", "图 B-9：多模态模型配置"),
        ("10-code-sandbox-trusted-local.png", "图 B-10：本地可信研究模式代码执行；不具备 Docker 容器隔离"),
    ):
        path = SCREENSHOTS / filename
        if path.exists():
            doc.add_picture(str(path), width=IMAGE_WIDTH)
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            add_inline(p, caption, size=9)

    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)
    print(f"created: {OUTPUT}")


if __name__ == "__main__":
    build()
