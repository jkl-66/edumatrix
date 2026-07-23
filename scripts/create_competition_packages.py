"""Build the three upload archives required by the competition portal.

The archives are deliberately separated:
  00015352作品.zip  - evaluator runnable package and one-click installer
  00015352源码.zip  - complete source package
  00015352介绍.zip  - Word/PDF/PPT/evidence and the optional demo video
"""

from __future__ import annotations

import hashlib
import json
import os
import re
import shutil
import sqlite3
import subprocess
import zipfile
from datetime import datetime, timezone
from pathlib import Path


ROOT = Path(__file__).resolve().parents[1]
OUTPUT = ROOT / "outputs"
PACKAGING = ROOT / "packaging"
BUILD = ROOT / "outputs" / "package_build"
RAW_REPOS = ROOT / "data" / "raw" / "github_repos"
RAW_IMAGE_SUFFIXES = {".png", ".jpg", ".jpeg", ".gif", ".svg", ".webp"}
RAW_MAX_IMAGE_BYTES = 512 * 1024

WORK_ZIP = OUTPUT / "00015352作品.zip"
SOURCE_ZIP = OUTPUT / "00015352源码.zip"
INFO_ZIP = OUTPUT / "00015352介绍.zip"

COMMON_ROOT_FILES = [
    ".env.example",
    ".dockerignore",
    "LICENSE",
    "NOTICE",
    "README.md",
    "requirements.txt",
    "pytest.ini",
    "run.py",
]

SOURCE_ROOT_FILES = COMMON_ROOT_FILES + ["Dockerfile", "docker-compose.yml", "start.bat"]
RUNTIME_SCRIPTS = [
    "db_heal.py",
    "runtime_preflight.py",
    "seed_students.py",
    "bootstrap_item_bank.py",
]

INFO_DOCS = [
    ("outputs/README.md", "README.md"),
    ("outputs/EduMatrix_评委阅读说明.md", "00_评委阅读说明.md"),
    ("outputs/EduMatrix_评委最终技术文档.docx", "01_EduMatrix_评委最终技术文档.docx"),
    ("outputs/EduMatrix_评委最终技术文档.pdf", "01_EduMatrix_评委最终技术文档.pdf"),
    ("outputs/EduMatrix_完整技术文档总稿.md", "02_EduMatrix_完整技术文档总稿.md"),
    ("outputs/EduMatrix_系统设计与实现方案.md", "03_EduMatrix_系统设计与实现方案.md"),
    ("outputs/EduMatrix_API与数据字典.md", "04_EduMatrix_API与数据字典.md"),
    ("outputs/EduMatrix_测试说明书.md", "05_EduMatrix_测试说明书.md"),
    ("outputs/EduMatrix_部署运维说明书.md", "06_EduMatrix_部署运维说明书.md"),
    ("outputs/EduMatrix_评委环境安装与复现备忘录.md", "07_EduMatrix_评委环境安装与复现备忘录.md"),
    ("outputs/EduMatrix_需求实现测试追踪矩阵.md", "08_EduMatrix_需求实现测试追踪矩阵.md"),
    ("outputs/EduMatrix_公开依据与引用清单.md", "09_EduMatrix_公开依据与引用清单.md"),
    ("outputs/EduMatrix_A3_评委阅读版_方向A.pptx", "10_EduMatrix_A3_评委阅读版_方向A.pptx"),
    ("outputs/EduMatrix_A3_评委阅读版_方向A.pdf", "10_EduMatrix_A3_评委阅读版_方向A.pdf"),
]


def sha256(path: Path) -> str:
    digest = hashlib.sha256()
    with path.open("rb") as handle:
        for block in iter(lambda: handle.read(1024 * 1024), b""):
            digest.update(block)
    return digest.hexdigest()


def add_file(archive: zipfile.ZipFile, source: Path, target: str, added: set[str]) -> bool:
    if not source.is_file():
        return False
    target = target.replace("\\", "/")
    if target in added:
        return True
    archive.write(source, target)
    added.add(target)
    return True


def safe_path(path: Path) -> bool:
    excluded_parts = {
        ".git", ".venv", "node_modules", "__pycache__", ".pytest_cache", ".mypy_cache",
        ".codex", ".agents", ".claude", ".cursor", ".gemini", ".grok", ".qwen",
        ".reasonix", ".serena", ".trae", ".ipynb_checkpoints", "_build", "build",
        "dist", "site", "html", "__macosx", "render", "recovery", "package_build",
    }
    if any(part in excluded_parts for part in path.parts):
        return False
    if path.suffix.lower() in {
        ".pyc", ".pyo", ".log", ".db", ".sqlite", ".wal", ".shm", ".pack", ".idx",
        ".rev", ".icloud", ".ds_store", ".doctree", ".inv", ".buildinfo", ".joblib",
        ".pkl", ".pickle", ".ckpt", ".lock", ".gz",
    }:
        return False
    if path.name.startswith("~$") or path.name.lower() in {".env", "proxy.log", "task.md", "test_member6_all_tasks.py"}:
        return False
    return True


def raw_path_allowed(path: Path) -> bool:
    if not safe_path(path):
        return False
    return not (
        path.suffix.lower() in RAW_IMAGE_SUFFIXES
        and path.stat().st_size > RAW_MAX_IMAGE_BYTES
    )


def add_tree(
    archive: zipfile.ZipFile,
    source: Path,
    prefix: str,
    added: set[str],
    path_allowed=safe_path,
) -> int:
    if not source.is_dir():
        return 0
    count = 0
    for path in sorted(source.rglob("*")):
        if not path.is_file() or not path_allowed(path):
            continue
        relative = path.relative_to(source).as_posix()
        if add_file(archive, path, f"{prefix.rstrip('/')}/{relative}", added):
            count += 1
    return count


def build_readme_docx(source: Path, target: Path, title: str) -> None:
    """Render a compact, printable Word version of a package README."""
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Pt
    from docx.oxml.ns import qn

    document = Document()
    section = document.sections[0]
    section.top_margin = Pt(50)
    section.bottom_margin = Pt(50)
    section.left_margin = Pt(58)
    section.right_margin = Pt(58)

    normal = document.styles["Normal"]
    normal.font.name = "等线"
    normal.font.size = Pt(10.5)
    normal._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), "等线")
    normal.paragraph_format.space_after = Pt(4)
    normal.paragraph_format.line_spacing = 1.12

    for style_name, size in (("Heading 1", 16), ("Heading 2", 13), ("Heading 3", 11)):
        style = document.styles[style_name]
        style.font.name = "等线"
        style.font.size = Pt(size)
        style.font.bold = True
        style._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), "等线")

    heading = document.add_paragraph()
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = heading.add_run(title)
    run.bold = True
    run.font.name = "等线"
    run.font.size = Pt(20)
    heading.paragraph_format.space_after = Pt(12)

    source_note = document.add_paragraph(f"文件来源：{source.name}")
    source_note.alignment = WD_ALIGN_PARAGRAPH.CENTER
    source_note.runs[0].italic = True
    source_note.runs[0].font.size = Pt(9)

    in_code = False
    for raw_line in source.read_text(encoding="utf-8", errors="replace").splitlines():
        line = raw_line.rstrip()
        stripped = line.strip()
        if stripped.startswith("```"):
            in_code = not in_code
            continue
        if in_code:
            paragraph = document.add_paragraph()
            paragraph.paragraph_format.left_indent = Pt(12)
            code_run = paragraph.add_run(line)
            code_run.font.name = "Consolas"
            code_run.font.size = Pt(9)
            continue
        if not stripped:
            document.add_paragraph().paragraph_format.space_after = Pt(1)
            continue
        if stripped == "---":
            document.add_paragraph("-" * 72)
            continue
        if stripped.startswith("#"):
            level = min(len(stripped) - len(stripped.lstrip("#")), 3)
            text = stripped[level:].strip()
            document.add_paragraph(text, style=f"Heading {level}")
            continue
        if stripped.startswith(">"):
            paragraph = document.add_paragraph(style="Intense Quote")
            paragraph.add_run(stripped.lstrip("> "))
            continue
        if re.match(r"^[-*+]\s+", stripped):
            document.add_paragraph(re.sub(r"^[-*+]\s+", "", stripped), style="List Bullet")
            continue
        if re.match(r"^\d+[.]\s+", stripped):
            document.add_paragraph(re.sub(r"^\d+[.]\s+", "", stripped), style="List Number")
            continue
        document.add_paragraph(stripped)

    document.core_properties.title = title
    document.core_properties.subject = "EduMatrix package README"
    target.parent.mkdir(parents=True, exist_ok=True)
    document.save(target)


def build_readme_documents() -> dict[str, Path]:
    BUILD.mkdir(parents=True, exist_ok=True)
    documents = {
        "project": BUILD / "README.docx",
        "work": BUILD / "README_作品运行说明.docx",
        "source": BUILD / "README_SOURCE_PACKAGE.docx",
        "info_overview": BUILD / "README_INFO_OVERVIEW.docx",
        "info": BUILD / "README_INFO_PACKAGE.docx",
    }
    build_readme_docx(ROOT / "README.md", documents["project"], "EduMatrix 项目 README")
    build_readme_docx(
        PACKAGING / "README_作品运行说明.txt",
        documents["work"],
        "EduMatrix 作品安装与运行说明",
    )
    build_readme_docx(
        PACKAGING / "README_源码包说明.txt",
        documents["source"],
        "EduMatrix 源码包说明",
    )
    build_readme_docx(
        OUTPUT / "README.md",
        documents["info_overview"],
        "EduMatrix 评委材料目录",
    )
    build_readme_docx(
        PACKAGING / "README_介绍包说明.txt",
        documents["info"],
        "EduMatrix 介绍材料包说明",
    )
    return documents


def build_raw_repo_manifest() -> Path:
    """Create a reviewer-facing inventory for the bundled public repo materials."""
    BUILD.mkdir(parents=True, exist_ok=True)
    source_metadata: dict[str, dict] = {}
    source_manifest = ROOT / "data" / "manifest" / "source_manifest.jsonl"
    if source_manifest.is_file():
        for line in source_manifest.read_text(encoding="utf-8", errors="replace").splitlines():
            try:
                record = json.loads(line)
            except json.JSONDecodeError:
                continue
            if record.get("key"):
                source_metadata[str(record["key"])] = record
    repositories: list[dict] = []
    total_files = 0
    total_bytes = 0
    if RAW_REPOS.is_dir():
        for repo in sorted(path for path in RAW_REPOS.iterdir() if path.is_dir()):
            selected_files: list[Path] = []
            excluded_images: list[dict] = []
            for path in sorted(repo.rglob("*")):
                if path.is_file() and safe_path(path) and not raw_path_allowed(path):
                    if path.suffix.lower() in RAW_IMAGE_SUFFIXES:
                        excluded_images.append({
                            "path": path.relative_to(RAW_REPOS).as_posix(),
                            "bytes": path.stat().st_size,
                        })
                if path.is_file() and raw_path_allowed(path):
                    selected_files.append(path)
            license_files = [
                path.relative_to(RAW_REPOS).as_posix()
                for path in selected_files
                if path.name.lower().startswith(("license", "copying", "notice"))
            ]
            bytes_count = sum(path.stat().st_size for path in selected_files)
            total_files += len(selected_files)
            total_bytes += bytes_count
            repositories.append({
                "repository": repo.name,
                "source_metadata": source_metadata.get(repo.name, {}),
                "selected_files": len(selected_files),
                "selected_bytes": bytes_count,
                "license_candidates": license_files,
                "excluded_oversized_images": excluded_images,
            })
    manifest = {
        "source_root": "data/raw/github_repos",
        "purpose": "预置课程知识库原始输入与可追溯素材",
        "selection": {
            "included": "仓库中的课程文档、Notebook、代码、数据、演示文稿、图像和许可证文件",
            "excluded": "Git 历史对象、构建产物、缓存、解释器缓存、私密配置、运行数据库及模型二进制缓存",
            "license_note": "许可证候选文件仅作来源索引；实际使用须遵守每个仓库随附的许可证及原始平台条款。",
        },
        "repositories": repositories,
        "totals": {"selected_files": total_files, "selected_bytes": total_bytes},
    }
    target = BUILD / "github_repos_package_manifest.json"
    target.write_text(json.dumps(manifest, ensure_ascii=False, indent=2), encoding="utf-8")
    return target


def lock_text() -> str:
    candidates = [ROOT / ".venv" / "Scripts" / "python.exe", Path(os.sys.executable)]
    for python in candidates:
        if not python.is_file():
            continue
        try:
            result = subprocess.run(
                [str(python), "-m", "pip", "freeze", "--all"],
                capture_output=True,
                text=True,
                encoding="utf-8",
                errors="replace",
                check=True,
            )
            return result.stdout
        except (OSError, subprocess.CalledProcessError):
            continue
    return "# pip freeze unavailable on the packaging host\n"


def build_seed_database() -> Path | None:
    source = ROOT / "edumatrix.db"
    if not source.is_file():
        return None
    BUILD.mkdir(parents=True, exist_ok=True)
    target = BUILD / "edumatrix.db"
    if target.exists():
        target.unlink()
    source_conn = sqlite3.connect(source)
    target_conn = sqlite3.connect(target)
    try:
        source_conn.backup(target_conn)
        tables = [row[0] for row in target_conn.execute(
            "select name from sqlite_master where type='table' and name not like 'sqlite_%'"
        )]
        for table in tables:
            columns = {row[1] for row in target_conn.execute(f'pragma table_info("{table}")')}
            quoted = '"' + table.replace('"', '""') + '"'
            if "student_id" in columns:
                if table == "knowledge_documents":
                    target_conn.execute(f"delete from {quoted} where student_id <> 'public'")
                else:
                    target_conn.execute(
                        f"delete from {quoted} where student_id not like 'stu-%' "
                        "and student_id not in ('public', 'demo-student')"
                    )
            if table == "users":
                target_conn.execute(
                    f"delete from {quoted} where username <> 'teacher' "
                    "and username <> 'demo-student' and username not like 'stu-%'"
                )
        target_conn.commit()
        target_conn.execute("vacuum")
    finally:
        source_conn.close()
        target_conn.close()
    return target


def evaluator_report() -> dict:
    source = OUTPUT / "e2e_no_docker" / "report.json"
    if not source.is_file():
        return {"result": "not_available"}
    raw = json.loads(source.read_text(encoding="utf-8"))
    checks = raw.get("checks", {})
    health = dict(checks.get("health", {}))
    sandbox = dict(checks.get("sandbox_status", {}))
    for key in ("llm_endpoint", "llm_api_key_configured", "llm_model", "user_api_key_provided"):
        health.pop(key, None)
    for key in ("python_executable", "multiprocessing_executable", "process_id"):
        sandbox.pop(key, None)
    return {
        "result": raw.get("result", "unknown"),
        "mode": raw.get("mode", "disabled"),
        "started_at": raw.get("started_at"),
        "finished_at": raw.get("finished_at"),
        "checks": {
            "health": health,
            "sandbox_status": sandbox,
            "chat_response": checks.get("chat_response", {}),
            "frontend_sandbox_guard": checks.get("frontend_sandbox_guard", {}),
            "knowledge_page": checks.get("knowledge_page", {}),
            "profile_analysis_page": checks.get("profile_analysis_page", {}),
            "settings_page": checks.get("settings_page", {}),
        },
    }


def video_files() -> list[Path]:
    candidates: list[Path] = []
    for directory in (ROOT / "submission_media", ROOT / "media"):
        if not directory.is_dir():
            continue
        candidates.extend(
            path for path in sorted(directory.rglob("*"))
            if path.is_file() and path.suffix.lower() in {".mp4", ".mov", ".mkv", ".webm"}
        )
    return candidates


def add_common_runtime(archive: zipfile.ZipFile, added: set[str]) -> int:
    count = 0
    for filename in COMMON_ROOT_FILES:
        count += int(add_file(archive, ROOT / filename, filename, added))
    for path in sorted(ROOT.glob("*.py")):
        count += int(add_file(archive, path, path.name, added))
    count += add_tree(archive, ROOT / "app", "app", added)
    for filename in RUNTIME_SCRIPTS:
        count += int(add_file(archive, ROOT / "scripts" / filename, f"scripts/{filename}", added))
    return count


def build_work_package(
    seed_db: Path | None,
    lock: str,
    readmes: dict[str, Path],
    raw_manifest: Path,
) -> None:
    added: set[str] = set()
    with zipfile.ZipFile(WORK_ZIP, "w", zipfile.ZIP_DEFLATED, compresslevel=6) as archive:
        add_common_runtime(archive, added)
        add_file(archive, ROOT / "frontend" / "dist" / "index.html", "frontend/dist/index.html", added)
        add_tree(archive, ROOT / "frontend" / "dist", "frontend/dist", added)
        for directory in ("datasets", "manifest", "patches"):
            add_tree(archive, ROOT / "data" / directory, f"data/{directory}", added)
        add_tree(archive, RAW_REPOS, "data/raw/github_repos", added, raw_path_allowed)
        add_file(archive, raw_manifest, "data/manifest/github_repos_package_manifest.json", added)
        for filename in ("poincare_projection.npy", "dkt_weights.pth"):
            add_file(archive, ROOT / "data" / filename, f"data/{filename}", added)
        if seed_db:
            add_file(archive, seed_db, "edumatrix.db", added)
        for filename in (
            "install_and_run.bat", "stop_services.bat", "verify_runtime.bat",
            ".env.runtime.template", "README_作品运行说明.txt",
        ):
            add_file(archive, PACKAGING / filename, filename, added)
        add_file(archive, readmes["project"], "README.docx", added)
        add_file(archive, readmes["work"], "README_作品运行说明.docx", added)
        archive.writestr("requirements-lock.txt", lock)
        archive.writestr("PACKAGE_MANIFEST.json", json.dumps({
            "package": "00015352作品.zip",
            "runtime": "Windows + Python 3.11",
            "frontend": "frontend/dist served by FastAPI on port 8000",
            "knowledge_source": "data/raw/github_repos (13 repositories, filtered public course materials)",
            "knowledge_manifest": "data/manifest/github_repos_package_manifest.json",
            "default_mode": "deterministic + sandbox disabled",
            "dependency_install": "online first-run pip install from requirements.txt",
            "created_at": datetime.now(timezone.utc).isoformat(),
        }, ensure_ascii=False, indent=2))


def build_source_package(lock: str, readmes: dict[str, Path]) -> None:
    added: set[str] = set()
    with zipfile.ZipFile(SOURCE_ZIP, "w", zipfile.ZIP_DEFLATED, compresslevel=6) as archive:
        for filename in SOURCE_ROOT_FILES:
            add_file(archive, ROOT / filename, filename, added)
        for path in sorted(ROOT.glob("*.py")):
            add_file(archive, path, path.name, added)
        for directory in ("app", "scripts", "tests", "docs", "packaging"):
            add_tree(archive, ROOT / directory, directory, added)
        add_tree(archive, ROOT / "frontend" / "src", "frontend/src", added)
        add_tree(archive, ROOT / "frontend" / "public", "frontend/public", added)
        for filename in ("package.json", "package-lock.json", "vite.config.js", "README.md"):
            add_file(archive, ROOT / "frontend" / filename, f"frontend/{filename}", added)
        # The source archive carries provenance metadata, not the evaluator's
        # course corpus. The runnable archive is the owner of those inputs.
        for directory in ("manifest",):
            add_tree(archive, ROOT / "data" / directory, f"data/{directory}", added)
        for filename in ("poincare_projection.npy", "dkt_weights.pth"):
            add_file(archive, ROOT / "data" / filename, f"data/{filename}", added)
        archive.writestr("requirements-lock.txt", lock)
        archive.writestr("README_SOURCE_PACKAGE.txt", (PACKAGING / "README_源码包说明.txt").read_text(encoding="utf-8"))
        add_file(archive, readmes["project"], "README.docx", added)
        add_file(archive, readmes["source"], "README_SOURCE_PACKAGE.docx", added)
        archive.writestr("PACKAGE_MANIFEST.json", json.dumps({
            "package": "00015352源码.zip",
            "contains": ["backend_source", "frontend_source", "tests", "source_metadata", "dependency_files"],
            "excludes": [".env", "knowledge_base_raw_materials", "course_datasets", "database_runtime_files", "frontend_dist", "node_modules", "virtual_environment", "private_logs"],
            "created_at": datetime.now(timezone.utc).isoformat(),
        }, ensure_ascii=False, indent=2))


def build_info_package() -> None:
    added: set[str] = set()
    with zipfile.ZipFile(INFO_ZIP, "w", zipfile.ZIP_DEFLATED, compresslevel=6) as archive:
        for source_name, target_name in INFO_DOCS:
            add_file(archive, ROOT / source_name, target_name, added)
        add_tree(archive, OUTPUT / "figures", "figures", added)
        # Keep only the image evidence. The raw capture report contains local
        # paths and process identifiers; evaluator_report.json below is the
        # sanitized reviewer-facing equivalent.
        screenshot_root = OUTPUT / "evaluator_screenshots"
        for path in sorted(screenshot_root.rglob("*")):
            if not path.is_file() or path.name == "report.json" or not safe_path(path):
                continue
            add_file(archive, path, f"screenshots/{path.relative_to(screenshot_root).as_posix()}", added)
        for path in sorted((ROOT / "render").glob("slide*.png")):
            add_file(archive, path, f"ppt_preview/{path.name}", added)
        add_file(archive, ROOT / "render" / "viewer.html", "ppt_preview/viewer.html", added)
        add_file(archive, OUTPUT / "runtime_security_matrix.json", "runtime_security_matrix.json", added)
        add_file(archive, OUTPUT / "trusted_local_smoke.json", "trusted_local_smoke.json", added)
        add_file(archive, BUILD / "README_INFO_OVERVIEW.docx", "README.docx", added)
        add_file(archive, BUILD / "README_INFO_PACKAGE.docx", "README_INFO_PACKAGE.docx", added)
        archive.writestr("evaluator_report.json", json.dumps(evaluator_report(), ensure_ascii=False, indent=2))
        videos = video_files()
        if videos:
            for video in videos:
                add_file(archive, video, f"video/{video.name}", added)
        else:
            archive.writestr(
                "video/VIDEO_PENDING.txt",
                "演示视频尚未由参赛队伍提供。收到视频后放入 submission_media/ 并重新运行 scripts/create_competition_packages.py。\n",
            )
        archive.writestr("README_INFO_PACKAGE.txt", (PACKAGING / "README_介绍包说明.txt").read_text(encoding="utf-8"))
        archive.writestr("PACKAGE_MANIFEST.json", json.dumps({
            "package": "00015352介绍.zip",
            "word_pages": 62,
            "ppt_pages": 20,
            "video_included": bool(videos),
            "created_at": datetime.now(timezone.utc).isoformat(),
        }, ensure_ascii=False, indent=2))


def main() -> None:
    OUTPUT.mkdir(parents=True, exist_ok=True)
    for path in (WORK_ZIP, SOURCE_ZIP, INFO_ZIP):
        if path.exists():
            path.unlink()
    seed_db = build_seed_database()
    lock = lock_text()
    readmes = build_readme_documents()
    raw_manifest = build_raw_repo_manifest()
    build_work_package(seed_db, lock, readmes, raw_manifest)
    build_source_package(lock, readmes)
    build_info_package()
    for path in (WORK_ZIP, SOURCE_ZIP, INFO_ZIP):
        size = path.stat().st_size
        if size >= 1_000_000_000:
            raise RuntimeError(f"{path.name} exceeds the 1 GB portal limit: {size} bytes")
        print(f"created {path} size={size} sha256={sha256(path)}")


if __name__ == "__main__":
    main()
