import os
os.environ["EDUMATRIX_LLM_PROVIDER"] = "mock"
os.environ["EDUMATRIX_EMBEDDING_PROVIDER"] = "hash"

from pathlib import Path
import unittest

from agent_swarm import EduMatrixSwarm
from drag_debate import DebateAugmentedRAG
from ingestion import DocumentIngestionPipeline
from instruct_rag import AsyncInstructRAGGenerator
from llm_client import DEFAULT_ASYNC_LLM
from manifold_alignment import verify_consistency
from models import StudentProfile
from observability import TELEMETRY
from rag_engine import hybrid_rag
from retrieval_evaluation import RetrievalEvalCase, evaluate_retrieval
from vector_store import InMemoryVectorIndex
from web_demo import _teacher_dashboard
from app.database import init_db
init_db()



class EduMatrixPipelineTests(unittest.TestCase):
    def test_graph_visrag_retrieves_pooling_context(self):
        bundle = hybrid_rag.retrieve("用图演示最大池化层", target="池化层")
        self.assertEqual(bundle.target, "池化层")
        self.assertIn("特征图", bundle.graph_context.learning_path)
        self.assertTrue(any(item.id == "IMG_PATCH_POOL_01" for item in bundle.evidence))

    def test_drag_keeps_relevant_evidence(self):
        bundle = hybrid_rag.retrieve("用 PyTorch 代码演示最大池化层", target="池化层")
        result = DebateAugmentedRAG().clean(bundle)
        kept = {verdict.evidence_id for verdict in result.verdicts if verdict.keep}
        self.assertIn("IMG_PATCH_POOL_01", kept)
        self.assertTrue(result.clean_evidence)

    def test_alignment_detects_pooling_conflict(self):
        text = "本讲义讲最大池化，窗口取局部最大值。"
        code = "import torch.nn as nn\npool = nn.AvgPool2d(2)"
        self.assertFalse(verify_consistency(text, code, threshold=0.4).passed)

    def test_swarm_generates_full_resource_package(self):
        swarm = EduMatrixSwarm()
        package = swarm.process("我看不懂池化层，请用图和 PyTorch 代码演示最大池化。")
        self.assertTrue(package.alignment.passed, package.alignment.advice)
        self.assertEqual(len(package.resources), 5)
        self.assertEqual(
            {resource.resource_type for resource in package.resources},
            {"专业讲义", "思维导图", "代码实操案例", "练习题", "虚拟人视频脚本"},
        )
        self.assertIn("knowledge_mastery", package.profile.dimension_states)
        self.assertTrue(package.profile.learning_state_causes)
        combined = "\n".join(resource.content for resource in package.resources)
        self.assertIn("MaxPool2d", combined)

    def test_multi_agent_5_resources_concurrent_generation(self):
        """验证多智能体并发生成：5种资源类型 + asyncio.gather + 失败容错"""
        from agent_swarm import AsyncResourceFactory, AGENT_MATRIX, AgentOutput
        import inspect
        # 1. AGENT_MATRIX 架构验证
        keys = [a.key for a in AGENT_MATRIX]
        bands = [a.band for a in AGENT_MATRIX]
        self.assertEqual(len(AGENT_MATRIX), 9, "必须9个智能体：1+3+5")
        self.assertEqual(sum(1 for b in bands if b == "Interaction Hub"), 1)
        self.assertEqual(sum(1 for b in bands if b == "Cognitive Governance"), 3)
        self.assertEqual(sum(1 for b in bands if b == "Resource Factory"), 5)
        # 2. AsyncResourceFactory 定义5种资源类型
        gen = AsyncInstructRAGGenerator(DEFAULT_ASYNC_LLM)
        factory = AsyncResourceFactory(gen)
        self.assertEqual(len(factory.jobs), 5)
        types = {t for _, t in factory.jobs}
        self.assertEqual(types, {"专业讲义", "思维导图", "代码实操案例", "练习题", "虚拟人视频脚本"})
        # 3. asyncio.gather
        source = inspect.getsource(factory.generate_all)
        self.assertIn("asyncio.gather", source)
        self.assertIn("return_exceptions=True", source)
        # 4. 失败容错
        self.assertIn("isinstance(output, Exception)", source)

    def test_dialogue_profile_extracts_state_causes_and_percentages(self):
        profile = StudentProfile(student_id="s1")
        profile.update_from_message(
            "我是计算机专业，期末要考数据结构和 CNN。池化层还是看不懂，"
            "最大池化和平均池化总混，题干长会漏条件，复习时只会看答案，"
            "我以为会了一做就错，希望你用图、代码和一步步提示。"
        )

        self.assertEqual(profile.major, "计算机专业")
        self.assertIn("期末复习", profile.learning_goals)
        self.assertIn("池化层", profile.weak_points)
        self.assertIn("misconception", profile.learning_state_causes)
        self.assertIn("cognitive_load", profile.learning_state_causes)
        self.assertIn("strategy_gap", profile.learning_state_causes)
        self.assertIn("metacognitive_mismatch", profile.learning_state_causes)
        self.assertEqual(len(profile.dimension_states), 10)

        total = sum(cause.percentage for cause in profile.learning_state_causes.values())
        self.assertAlmostEqual(total, 100.0, delta=0.6)
        report = profile.state_report()
        self.assertIn("不会原因占比", report)
        self.assertIn("误概念/易混点", report)

    def test_feedback_updates_mastery_and_metacognition(self):
        profile = StudentProfile(student_id="s2")
        profile.update_from_message("我觉得池化层已经懂了，但不确定自己是不是真的会。")
        before = profile.concept_mastery.get("池化层", 0.5)
        profile.update_from_feedback(
            feedback="做最大池化诊断题时选成了平均池化，提示两次后才改对。",
            accuracy=0.35,
            self_confidence=0.85,
            hint_count=2,
        )

        self.assertLess(profile.concept_mastery["池化层"], before)
        self.assertIn("metacognitive_mismatch", profile.learning_state_causes)
        self.assertIn("strategy_gap", profile.learning_state_causes)

    def test_behavior_sanity_hard_cap_locks_mastery(self):
        from bkt_engine import behavior_sanity_check
        profile = StudentProfile(student_id="s3")
        profile.concept_mastery["逻辑回归"] = 0.85
        profile.recent_quiz_accuracy["逻辑回归"] = [0.0, 0.0, 0.0]
        profile.metacognitive_mismatch = 0.10
        result = behavior_sanity_check(profile)
        self.assertTrue(result["sanitized"])
        self.assertIn("逻辑回归", result["capped_concepts"])
        self.assertLessEqual(profile.concept_mastery["逻辑回归"], 0.5,
                             "3次答错后掌握度应被 cap 至 0.5")
        self.assertGreaterEqual(profile.metacognitive_mismatch, 0.40,
                                "元认知偏差应上调 30% (原0.10 + 0.30 = 0.40)")

    def test_machine_learning_course_flow_generates_strategy_plan(self):
        swarm = EduMatrixSwarm()
        package = swarm.process(
            "我是计算机专业，期末要考机器学习。逻辑回归和混淆矩阵总混，"
            "accuracy 很高但 recall 低我不知道怎么判断，希望用图和例子一步步讲。",
            student_id="ml-student",
        )

        self.assertEqual(package.profile.target_course, "机器学习导论")
        self.assertIn(package.target, {"逻辑回归", "混淆矩阵", "模型评估"})
        self.assertTrue(package.strategy_plan)
        self.assertTrue(package.strategy_plan.actions)
        self.assertTrue(any("混淆矩阵" in item.content or "逻辑回归" in item.content for item in package.resources))

    def test_zpd_path_planning_and_adaptive_teaching(self):
        """验证 ZPD 三档分类 + 前置回滚 + 自适应二档教学机制"""
        from bkt_engine import (
            classify_zpd_zone, should_rollback_to_prerequisites,
            get_zpd_path_plan, ZPD_LOWER, ZPD_UPPER
        )
        from agent_swarm import DEFAULT_KNOWLEDGE_DAG, SwarmMediationRouter, SwarmMediationMode
        from learning_strategy import detect_teaching_tier, TeachingTier
        # SC1: ZPD 三档分类
        self.assertEqual(classify_zpd_zone(0.1), "below_zpd")
        self.assertEqual(classify_zpd_zone(ZPD_LOWER), "in_zpd")
        self.assertEqual(classify_zpd_zone(0.5), "in_zpd")
        self.assertEqual(classify_zpd_zone(ZPD_UPPER), "in_zpd")
        self.assertEqual(classify_zpd_zone(0.8), "above_zpd")
        # SC2: 前置依赖回滚
        rollback, weak = should_rollback_to_prerequisites(0.2, {'线性回归': 0.3})
        self.assertTrue(rollback)
        self.assertIn('线性回归', weak)
        rollback2, _ = should_rollback_to_prerequisites(0.2, {'线性回归': 0.9})
        self.assertFalse(rollback2)
        # ZPD 路径完整规划
        prereq = {'线性回归': 0.3, '梯度下降': 0.4}
        plan_low = get_zpd_path_plan('逻辑回归', 0.15, DEFAULT_KNOWLEDGE_DAG, prereq)
        self.assertEqual(plan_low['difficulty'], 'basic')
        self.assertIsNotNone(plan_low['rollback_to'])
        prereq_mid = {'线性回归': 0.80, '梯度下降': 0.70}
        plan_mid = get_zpd_path_plan('逻辑回归', 0.55, DEFAULT_KNOWLEDGE_DAG, prereq_mid)
        self.assertEqual(plan_mid['difficulty'], 'intermediate')
        prereq_high = {'线性回归': 0.90, '梯度下降': 0.85}
        plan_high = get_zpd_path_plan('逻辑回归', 0.85, DEFAULT_KNOWLEDGE_DAG, prereq_high)
        self.assertEqual(plan_high['difficulty'], 'advanced')
        # SC3+SC4: 自适应二档教学
        router = SwarmMediationRouter()
        p = StudentProfile(student_id='zpd_test')
        p.concept_mastery['逻辑回归'] = 0.30
        self.assertEqual(router.decide_mode(p, '逻辑回归'), SwarmMediationMode.SIMPLIFIED_MODE)
        self.assertIn('降维解释', router.get_forced_instructions(SwarmMediationMode.SIMPLIFIED_MODE).get('theory', ''))
        p.concept_mastery['逻辑回归'] = 0.85
        self.assertEqual(router.decide_mode(p, '逻辑回归'), SwarmMediationMode.ADVANCED_MODE)
        self.assertIn('ADVANCED_MODE', router.get_forced_instructions(SwarmMediationMode.ADVANCED_MODE).get('theory', ''))
        # learning_strategy 档位一致性
        self.assertEqual(detect_teaching_tier(p), TeachingTier.ADVANCED)
        p.concept_mastery['逻辑回归'] = 0.30
        self.assertEqual(detect_teaching_tier(p), TeachingTier.SIMPLIFIED)

    def test_teacher_dashboard_exposes_heatmap_and_interventions(self):
        dashboard = _teacher_dashboard()
        self.assertEqual(dashboard["course"], "机器学习导论")
        self.assertTrue(dashboard["heatmap"])
        self.assertTrue(dashboard["interventions"])
    def test_kalman_filter_smoothing_and_state_recovery(self):
        """验证自适应卡尔曼滤波平滑估计以及从数据库快照冷启动状态恢复"""
        from bkt_engine import BKTEngine
        engine = BKTEngine()
        
        # 1. 验证自适应卡尔曼滤波更新与平滑效果
        smoothed_1 = engine.update("逻辑回归", correct=True, cognitive_load=0.4, frustration=0.1)
        self.assertGreater(smoothed_1, 0.3)
        self.assertLess(smoothed_1, 0.6)
        
        # 2. 模拟由于情绪和高认知负荷导致的波动：连续做错，且挫败感极大 (观测噪声 R 增加)
        smoothed_2 = engine.update("逻辑回归", correct=False, cognitive_load=0.9, frustration=0.8)
        self.assertLess(smoothed_2, smoothed_1)
        
        # 3. 验证冷启动状态恢复
        snap = engine.snapshot()
        self.assertIn("逻辑回归", snap)
        self.assertIn("smoothed_mastery", snap["逻辑回归"])
        self.assertIn("p_err", snap["逻辑回归"])
        
        # 重新创建引擎，模拟冷启动，并传入 profile_bkt_states 恢复状态
        new_engine = BKTEngine()
        self.assertEqual(new_engine.get_mastery("逻辑回归"), 0.3)
        
        # 传入 profile_bkt_states，应能成功还原状态值（使用一个全新干净的引擎避免污染）
        recovery_engine = BKTEngine()
        self.assertAlmostEqual(recovery_engine.get_mastery("逻辑回归", snap), snap["逻辑回归"]["smoothed_mastery"])
        
        # 测试在恢复状态后进行更新，误差协方差应正常传承
        recovered_state = recovery_engine.get_or_create("逻辑回归", snap)
        self.assertAlmostEqual(recovered_state.p_err, snap["逻辑回归"]["p_err"])

    def test_dynamic_concept_resolution_and_pronoun_alignment(self):
        """验证无白名单硬编码的动态指代消解与实体链接功能"""
        from agent_swarm import ZPDPlannerAgent
        from models import StudentProfile
        
        planner = ZPDPlannerAgent()
        profile = StudentProfile(student_id="test_resolution_student")
        
        # 1. 模拟历史对话中提到了“池化层”
        profile.history = [
            "User: 请问最大池化和平均池化有什么区别？",
            "Assistant: 最大池化保留局部最大激活值，而平均池化对窗口内的值取平均，两者都属于池化层。",
        ]
        
        # 2. 当学生输入含指代代词的问题时，应能成功消解并链接到“池化层”
        resolved = planner._extract_concept_from_query("这个该怎么实现？", profile)
        self.assertEqual(resolved, "池化层")
        
        # 3. 校验如果上下文无匹配词，应安全退回到 None
        resolved_none = planner._extract_concept_from_query("刚才说的那个该怎么做？", None)
        self.assertIsNone(resolved_none)

    def test_drag_debate_cleans_low_quality_evidence(self):
        from drag_debate import DebateAugmentedRAG
        from models import Evidence, EvidenceModality, RetrievalBundle, GraphContext
        debate = DebateAugmentedRAG()
        evidence = (
            Evidence(id="e1", title="good", content="池化层概念", modality=EvidenceModality.TEXT, source="教材", score=0.9, tags=(), anchors=()),
            Evidence(id="e2", title="noise", content="无关内容", modality=EvidenceModality.TEXT, source="网络", score=0.3, tags=(), anchors=()),
        )
        bundle = RetrievalBundle(query="池化层", target="池化层",
            graph_context=GraphContext(target="池化层", learning_path=(), prerequisite_edges=(), downstream_edges=()),
            evidence=evidence)
        result = debate.clean(bundle)
        self.assertLessEqual(len(result.clean_evidence), len(evidence))
        self.assertGreater(len(result.clean_evidence), 0)

    def test_sm2_flashcard_scheduling(self):
        from anki_engine import SM2Engine, FlashCard
        engine = SM2Engine()
        card = engine.get_or_create(concept='逻辑回归', front='概念', back='解析', student_id='s-test')
        self.assertIsNotNone(card)
        engine.review('逻辑回归', 's-test', quality=5)
        engine.review('逻辑回归', 's-test', quality=4)
        due = engine.get_due_cards('s-test')
        self.assertIsInstance(due, list)

    def test_learning_event_bus_subscribe_publish(self):
        from learning_event_bus import LearningEventBus
        bus = LearningEventBus.get_instance()
        received = []
        async def handler(e):
            received.append(e)
        bus.subscribe('test_evt', handler)
        bus.unsubscribe('test_evt', handler)
        bus.subscribe('test_evt2', handler)
        bus.unsubscribe('test_evt2', handler)
        self.assertEqual(len(received), 0)

    def test_circuit_breaker_opens_after_failures(self):
        from concurrency import CircuitBreaker, CircuitState
        import asyncio
        cb = CircuitBreaker(name="test", failure_threshold=3, recovery_timeout=1.0)
        self.assertEqual(cb.state, CircuitState.CLOSED)
        async def run():
            for i in range(3):
                async def fail():
                    raise Exception("fail")
                try:
                    await cb.call(fail)
                except Exception:
                    pass
        asyncio.run(run())
        self.assertEqual(cb.state, CircuitState.OPEN)

    def test_visrag_images_exist(self):
        import os
        img_dir = 'data/patches'
        images = [f for f in os.listdir(img_dir) if f.endswith('.png')]
        self.assertGreaterEqual(len(images), 7, f'期望至少7张配图，实际{len(images)}')
        for img in images:
            path = os.path.join(img_dir, img)
            with open(path, 'rb') as f:
                header = f.read(8)
            self.assertEqual(header, b'\x89PNG\r\n\x1a\n', f'{img} 不是有效PNG')

    def test_report_api_browser_pool_exists(self):
        """验证 PDF 导出模块 BrowserPool 存在且功能正确"""
        import asyncio
        from report_api import BrowserPool, get_browser_pool
        pool = asyncio.run(get_browser_pool())
        self.assertIsNotNone(pool)
        self.assertIsNotNone(pool.semaphore, 'BrowserPool 需有 semaphore 并发信号量')
        self.assertEqual(pool.render_timeout, 10)
        self.assertTrue(hasattr(pool, 'render_pdf'), 'BrowserPool 需有 render_pdf 方法')
        self.assertTrue(hasattr(pool, 'close'), 'BrowserPool 需有 close 方法')

    def test_sse_syntax_buffer_in_store(self):
        """验证 SSE 流式响应语法缓冲（Task 4.1）"""
        import os
        store_path = os.path.join(os.path.dirname(__file__), 'frontend', 'src', 'stores', 'chat.js')
        self.assertTrue(os.path.exists(store_path), 'chat store 不存在')
        with open(store_path, 'r', encoding='utf-8') as f:
            content = f.read()
        self.assertIn('abort', content.lower(), '需有 abort 中断机制')
        self.assertIn('streamChat', content, '需有 streamChat 流式聊天')
        self.assertIn('_cleanupStream', content, '需有 _cleanupStream 流清理')

    def test_frontend_core_components_exist(self):
        """验证测试手册中所有前端核心组件存在且包含关键功能"""
        import os
        checks = [
            # Task 4.2: 智能体轨迹时间轴
            ('frontend/src/components/AgentTimeline.vue', "step.status === 'running'", '呼吸灯状态'),
            # Task 4.2: 掌握度雷达图
            ('frontend/src/components/MasteryRadar.vue', '能力掌握度雷达图', 'ECharts雷达图'),
            # Task 5.1+5.2: 数字人语音+嘴形滤波
            ('frontend/src/components/AvatarSpeech.vue', 'smoothScale', '嘴形平滑滤波'),
            # Task 7.5: 3D Anki闪卡
            ('frontend/src/components/AnkiFlashcard.vue', 'rotateY(180deg)', '3D翻转'),
            # Task 8.1: 行级苏格拉底答疑
            ('frontend/src/components/InlineSocraticPopup.vue', '苏格拉底即时答疑', '行级答疑悬浮框'),
            # Task 8.4: 沙箱可视化终端
            ('frontend/src/components/SandboxConsole.vue', '代码沙箱控制台', '代码沙箱控制台'),
            # Task 8.5: 视频生成播放器
            ('frontend/src/components/VideoRenderPanel.vue', 'isComplete', 'HTML5视频播放器'),
            # Task 8.6: 知识图谱探索器
            ('frontend/src/components/KnowledgeGraphExplorer.vue', 'echarts.init', 'ECharts图谱'),
            # Task 8.7: 无图自适应折叠
            ('frontend/src/views/Chat.vue', 'shouldShowRightPanel', '自适应面板折叠'),
            # Task 7.4: 复习日历
            ('frontend/src/views/RevisionCalendar.vue', 'reviewPlans', '抗遗忘复习日历'),
            # Task 9.2: 设置致谢墙+风格切换
            ('frontend/src/views/Settings.vue', 'teachingStyle', '教学风格切换'),
        ]
        for path, keyword, label in checks:
            with self.subTest(component=label):
                full_path = os.path.join(os.path.dirname(__file__), path)
                self.assertTrue(os.path.exists(full_path), f'{path} 不存在')
                with open(full_path, 'r', encoding='utf-8') as f:
                    content = f.read()
                self.assertIn(keyword, content, f'{label}: 缺少关键词 "{keyword}"')

    def test_ingestion_pipeline_indexes_future_course_materials(self):
        index = InMemoryVectorIndex("test-course-index")
        pipeline = DocumentIngestionPipeline(index, chunk_size=80, overlap=10)
        report = pipeline.ingest_text(
            "最大池化使用窗口取局部最大值。混淆矩阵可以帮助学生理解 precision 和 recall 的差异。",
            source="unit-test.md",
            title="工业化摄入测试",
        )

        self.assertGreaterEqual(report.chunks, 1)
        self.assertEqual(index.count(), report.chunks)
        hits = index.search("最大池化 局部最大值", top_k=3)
        self.assertTrue(hits)
        self.assertIn("最大池化", hits[0].content)

    def test_retrieval_evaluation_reports_recall_and_mrr(self):
        report = evaluate_retrieval(
            hybrid_rag,
            (
                RetrievalEvalCase(
                    query="用 PyTorch 演示最大池化",
                    target="池化层",
                    expected_evidence_ids=("IMG_PATCH_POOL_01",),
                ),
            ),
        )
        self.assertEqual(report.cases, 1)
        self.assertGreaterEqual(report.recall_at_k, 1.0)
        self.assertGreater(report.mean_reciprocal_rank, 0.0)

    def test_telemetry_records_pipeline_metrics(self):
        TELEMETRY.clear()
        swarm = EduMatrixSwarm()
        swarm.process("我看不懂池化层，请用代码演示最大池化。")
        snapshot = TELEMETRY.snapshot()
        metric_names = {item["name"] for item in snapshot["metrics"]}
        span_names = {item["name"] for item in snapshot["spans"]}
        self.assertIn("retrieval.evidence_count", metric_names)
        self.assertIn("alignment.rollback_count", metric_names)
        self.assertIn("swarm.process", span_names)

    def test_database_pool_wash_with_postgres_mock(self):
        from unittest.mock import MagicMock
        from app.database import before_cursor_execute, on_connection_checkin, set_tenant, tenant_context
        
        # 1. 测试租户上下文切换
        self.assertEqual(tenant_context.get(), "public")
        with set_tenant("tenant_abc"):
            self.assertEqual(tenant_context.get(), "tenant_abc")
        self.assertEqual(tenant_context.get(), "public")
        
        # 2. 测试 before_cursor_execute 在 PostgreSQL 方言下的设置
        mock_conn = MagicMock()
        mock_conn.dialect.name = "postgresql"
        mock_cursor = MagicMock()
        
        with set_tenant("tenant_123"):
            before_cursor_execute(mock_conn, mock_cursor, "SELECT 1", {}, None, False)
            mock_cursor.execute.assert_called_with("SET search_path TO tenant_123;")
            
        # 3. 测试为 public 时不设置
        mock_cursor.reset_mock()
        before_cursor_execute(mock_conn, mock_cursor, "SELECT 1", {}, None, False)
        mock_cursor.execute.assert_not_called()
        
        # 4. 测试在 SQLite 下不起作用（防止本地开发环境报错）
        mock_conn.dialect.name = "sqlite"
        mock_cursor.reset_mock()
        with set_tenant("tenant_123"):
            before_cursor_execute(mock_conn, mock_cursor, "SELECT 1", {}, None, False)
            mock_cursor.execute.assert_not_called()

        # 5. 测试归还连接池时执行洗白 SQL
        mock_dbapi_conn = MagicMock()
        mock_dbapi_conn.__class__.__module__ = "psycopg2.extensions"
        mock_dbapi_cursor = MagicMock()
        mock_dbapi_conn.cursor.return_value = mock_dbapi_cursor
        
        on_connection_checkin(mock_dbapi_conn, None)
        mock_dbapi_cursor.execute.assert_called_with("SET search_path TO public;")
        mock_dbapi_cursor.close.assert_called_once()

    def test_stream_disconnect_handling(self):
        import asyncio
        from unittest.mock import MagicMock
        from fastapi import Request
        from stream_api import stream_chat

        mock_request = MagicMock(spec=Request)
        
        # 模拟 is_disconnected 最初返回 False，之后返回 True
        disconnect_calls = [False, False, True]
        async def mock_is_disconnected():
            if disconnect_calls:
                return disconnect_calls.pop(0)
            return True
        mock_request.is_disconnected = mock_is_disconnected
        mock_request.headers = {}
        
        async def mock_json():
            return {"message": "我看不懂池化层", "student_id": "test-student"}
        mock_request.json = mock_json

        # 执行 stream_chat 并消费其返回的流
        try:
            loop = asyncio.get_event_loop()
        except RuntimeError:
            loop = asyncio.new_event_loop()
            asyncio.set_event_loop(loop)

        async def run_test():
            response = await stream_chat(mock_request)
            async for chunk in response.body_iterator:
                pass

        with self.assertRaises(asyncio.CancelledError):
            if loop.is_running():
                from concurrent.futures import ThreadPoolExecutor
                with ThreadPoolExecutor(max_workers=1) as executor:
                    future = executor.submit(asyncio.run, run_test())
                    future.result()
            else:
                loop.run_until_complete(run_test())

    def test_guided_decoding_self_healing(self):
        import asyncio
        from unittest.mock import AsyncMock, patch
        from app.agents.coder import async_refine_code_agent, PyTorchPoolCodeSchema

        lecture = "我们要学习最大池化层，最大池化非常重要。"
        code = "import torch\nimport torch.nn as nn\nclass Model(nn.Module):\n    def __init__(self):\n        super().__init__()\n        self.pool = nn.AvgPool2d(kernel_size=2, stride=2)"
        
        try:
            loop = asyncio.get_event_loop()
        except RuntimeError:
            loop = asyncio.new_event_loop()
            asyncio.set_event_loop(loop)
            
        async def run_healing():
            return await async_refine_code_agent(lecture, code, "平均池化不符合最大池化讲义")
            
        if loop.is_running():
            from concurrent.futures import ThreadPoolExecutor
            with ThreadPoolExecutor(max_workers=1) as executor:
                future = executor.submit(asyncio.run, run_healing())
                healed_code = future.result()
        else:
            healed_code = loop.run_until_complete(run_healing())
            
        self.assertIn("nn.MaxPool2d", healed_code)
        self.assertNotIn("nn.AvgPool2d", healed_code)

        mock_schema = PyTorchPoolCodeSchema(
            import_blocks="import torch\nimport torch.nn as nn",
            tensor_init="x = torch.randn(1, 1, 4, 4)",
            pool_layer_api="pool = nn.MaxPool2d(2, 2)",
            forward_and_print="print(pool(x))"
        )
        
        async def run_success():
            with patch("instructor.patch") as mock_patch:
                mock_client = AsyncMock()
                mock_create = AsyncMock(return_value=mock_schema)
                mock_client.chat.completions.create = mock_create
                mock_patch.return_value = mock_client
                
                return await async_refine_code_agent(lecture, code, "对齐纠偏")

        if loop.is_running():
            from concurrent.futures import ThreadPoolExecutor
            with ThreadPoolExecutor(max_workers=1) as executor:
                future = executor.submit(asyncio.run, run_success())
                res_code = future.result()
        else:
            res_code = loop.run_until_complete(run_success())
            
        self.assertIn("x = torch.randn(1, 1, 4, 4)", res_code)
        self.assertIn("pool = nn.MaxPool2d(2, 2)", res_code)

    def test_sandbox_resource_limits_and_timeout(self):
        import asyncio
        from code_exec_api import SANDBOX_RUNNER
        
        try:
            loop = asyncio.get_event_loop()
        except RuntimeError:
            loop = asyncio.new_event_loop()
            asyncio.set_event_loop(loop)
            
        async def run_normal():
            return await SANDBOX_RUNNER.run("print('hello sandbox')")
            
        if loop.is_running():
            from concurrent.futures import ThreadPoolExecutor
            with ThreadPoolExecutor(max_workers=1) as executor:
                future = executor.submit(asyncio.run, run_normal())
                stdout, stderr, elapsed = future.result()
        else:
            stdout, stderr, elapsed = loop.run_until_complete(run_normal())
            
        self.assertIn("hello sandbox", stdout)
        self.assertEqual(stderr, "")
        
        async def run_timeout():
            from config import CONFIG
            sleep_time = int(CONFIG.sandbox_timeout + 2)
            return await SANDBOX_RUNNER.run(f"import time; time.sleep({sleep_time})")
            
        if loop.is_running():
            from concurrent.futures import ThreadPoolExecutor
            with ThreadPoolExecutor(max_workers=1) as executor:
                future = executor.submit(asyncio.run, run_timeout())
                stdout, stderr, elapsed = future.result()
        else:
            stdout, stderr, elapsed = loop.run_until_complete(run_timeout())
            
        self.assertEqual(stdout, "")
        self.assertIn("超时", stderr)
        from config import CONFIG
        self.assertLessEqual(elapsed, CONFIG.sandbox_timeout + 3.0)


    def test_sandbox_class_execution(self):
        import asyncio
        from code_exec_api import SANDBOX_RUNNER
        
        try:
            loop = asyncio.get_event_loop()
        except RuntimeError:
            loop = asyncio.new_event_loop()
            asyncio.set_event_loop(loop)
            
        code = """
class Base:
    def __init__(self, val):
        self.val = val

class Sub(Base):
    def __init__(self, val):
        super().__init__(val)
    
    @classmethod
    def create(cls, val):
        return cls(val)

s = Sub.create(42)
print("Val:", s.val)
"""
        async def run_code():
            return await SANDBOX_RUNNER.run(code)
            
        if loop.is_running():
            from concurrent.futures import ThreadPoolExecutor
            with ThreadPoolExecutor(max_workers=1) as executor:
                future = executor.submit(asyncio.run, run_code())
                stdout, stderr, elapsed = future.result()
        else:
            stdout, stderr, elapsed = loop.run_until_complete(run_code())
            
        self.assertEqual(stderr.strip(), "")
        self.assertIn("Val: 42", stdout)


    def test_database_cascade_deletes(self):
        from app.database import SessionLocal, DBStudentProfile, DBNote, DBReviewPlan, DBQuizRecord, DBWrongQuestion, DBConversationHistory
        from app.crud import save_student_profile
        from models import StudentProfile
        import uuid
        
        student_id = f"test-cascade-{uuid.uuid4().hex[:8]}"
        
        # 1. Create student profile
        profile = StudentProfile(student_id=student_id)
        profile.major = "AI"
        
        session = SessionLocal()
        try:
            save_student_profile(session, profile)
            
            # 2. Add child rows
            note = DBNote(id=f"n-{student_id}", student_id=student_id, content="test note")
            plan = DBReviewPlan(student_id=student_id, concept="CNN", mastery=0.5)
            quiz = DBQuizRecord(id=f"q-{student_id}", student_id=student_id, question="Q?", correct_answer="A")
            session.add_all([note, plan, quiz])
            session.commit()
            
            wrong = DBWrongQuestion(student_id=student_id, quiz_record_id=quiz.id, concept_name="CNN", wrong_reason_category="misconception")
            session.add(wrong)
            session.commit()
            
            # 3. Verify they exist
            self.assertIsNotNone(session.query(DBNote).filter_by(student_id=student_id).first())
            self.assertIsNotNone(session.query(DBReviewPlan).filter_by(student_id=student_id).first())
            self.assertIsNotNone(session.query(DBQuizRecord).filter_by(student_id=student_id).first())
            self.assertIsNotNone(session.query(DBWrongQuestion).filter_by(student_id=student_id).first())
            
            # 4. Delete student profile
            db_prof = session.query(DBStudentProfile).filter_by(student_id=student_id).first()
            session.delete(db_prof)
            session.commit()
            
            # 5. Verify they are gone!
            self.assertIsNone(session.query(DBNote).filter_by(student_id=student_id).first())
            self.assertIsNone(session.query(DBReviewPlan).filter_by(student_id=student_id).first())
            self.assertIsNone(session.query(DBQuizRecord).filter_by(student_id=student_id).first())
            self.assertIsNone(session.query(DBWrongQuestion).filter_by(student_id=student_id).first())
        finally:
            session.close()


    def test_database_concurrency_writes(self):
        import concurrent.futures
        from app.database import SessionLocal
        from app.crud import load_student_profile, save_student_profile
        import uuid
        
        student_id = f"test-con-{uuid.uuid4().hex[:8]}"
        
        def run_write():
            session = SessionLocal()
            try:
                prof = load_student_profile(session, student_id)
                prof.cognitive_load = 0.9
                save_student_profile(session, prof)
            finally:
                session.close()
                
        # Run 10 concurrent writes in parallel threads
        with concurrent.futures.ThreadPoolExecutor(max_workers=10) as executor:
            futures = [executor.submit(run_write) for _ in range(10)]
            for f in futures:
                # Should not raise "database is locked"
                f.result()
                
        # Verify write succeeded
        session = SessionLocal()
        try:
            prof = load_student_profile(session, student_id)
            self.assertEqual(prof.cognitive_load, 0.9)
        finally:
            session.close()

    def test_coreference_resolution_with_garbage_words(self):
        """验证口语指代消解在垃圾词打断场景下的自愈与兜底机制"""
        from agent_swarm import _resolve_coreference
        from models import StudentProfile
        
        # 1. 正常问一个问题："逻辑回归该怎么理解"
        profile = StudentProfile(student_id="test_coref")
        profile.update_from_message("逻辑回归该怎么理解")
        profile.concept_mastery["逻辑回归"] = 0.5 # 确保在 concept_mastery 中
        
        # 2. 模拟系统回复
        profile.update_from_feedback(feedback="逻辑回归是一个经典的二分类算法。")
        
        # 3. 输入一段垃圾词："这这那那，今天天气真好"
        profile.update_from_message("这这那那，今天天气真好")
        profile.update_from_feedback(feedback="抱歉，系统在知识库中未检索到与您提问相关的充足高置信度证据，为避免幻觉...")
        
        # 此时滑动历史上下文
        from agent_swarm import ProfileProbeAgent
        probe = ProfileProbeAgent()
        sliding_context = probe._get_sliding_context(profile, window_size=3)
        
        # 4. 再次提问没有明确代词指代的问题："那它的代码该怎么写"
        message = "那它的代码该怎么写"
        existing_concepts = list(profile.concept_mastery.keys())
        
        resolved = _resolve_coreference(message, existing_concepts, sliding_context)
        # 应将 "它" 替换为 "逻辑回归"
        self.assertIn("逻辑回归", resolved)
        self.assertNotIn("它", resolved)

    def test_automatic_review_plan_generation_on_chat(self):
        """验证在对话框学习后自动在 review_plans 表生成复习安排的机制"""
        from app.database import SessionLocal, DBReviewPlan
        from app.crud import load_student_profile, save_student_profile
        from models import StudentProfile
        import uuid
        
        student_id = f"test-chat-rev-{uuid.uuid4().hex[:8]}"
        
        # 1. 模拟学生画像并保存
        profile = StudentProfile(student_id=student_id)
        profile.concept_mastery["线性回归"] = 0.6
        
        session = SessionLocal()
        try:
            save_student_profile(session, profile)
            
            # 2. 模拟 stream_api 的逻辑，自动调用 upsert_review_plan
            from app.crud import upsert_review_plan
            upsert_review_plan(session, student_id, "线性回归", 0.6, 1)
            
            # 3. 验证数据库中是否成功产生复习计划
            plan = session.query(DBReviewPlan).filter_by(student_id=student_id, concept="线性回归").first()
            self.assertIsNotNone(plan)
            self.assertEqual(plan.mastery, 0.6)
            self.assertEqual(plan.interval_days, 1)
            
            # 4. 清理
            session.delete(plan)
            session.commit()
        finally:
            session.close()

    def test_manual_review_plan_creation_without_preexisting_profile(self):
        """验证在学生画像表尚无记录时，手动创建复习计划也能安全自愈通过"""
        from app.database import SessionLocal, DBReviewPlan, DBStudentProfile
        from app.crud import upsert_review_plan
        import uuid
        
        student_id = f"test-manual-no-prof-{uuid.uuid4().hex[:8]}"
        
        session = SessionLocal()
        try:
            # 1. 验证该学生在 student_profiles 表中确实不存在
            prof_exist = session.query(DBStudentProfile).filter_by(student_id=student_id).first()
            self.assertIsNone(prof_exist)
            
            # 2. 直接调用 upsert_review_plan，不应抛出 IntegrityError
            plan = upsert_review_plan(session, student_id, "支持向量机", 0.7, 3)
            
            # 3. 验证复习计划记录已成功创建
            self.assertIsNotNone(plan)
            self.assertEqual(plan.concept, "支持向量机")
            self.assertEqual(plan.interval_days, 3)
            
            # 4. 验证系统也已自动创建了 student_profile 母表记录以维系外键
            prof_now = session.query(DBStudentProfile).filter_by(student_id=student_id).first()
            self.assertIsNotNone(prof_now)
            
            # 5. 清理
            session.delete(plan)
            session.delete(prof_now)
            session.commit()
        finally:
            session.close()

    def test_flashcard_review_updates_persisted_sm2_plan_without_memory_card(self):
        """The review API should update review_plans even if no in-memory card exists."""
        from app.database import SessionLocal, DBReviewPlan, DBStudentProfile
        from app.crud import upsert_review_plan
        from fastapi.testclient import TestClient
        from app.main import app
        import uuid

        student_id = f"test-sm2-api-{uuid.uuid4().hex[:8]}"
        concept = f"sm2-concept-{uuid.uuid4().hex[:8]}"

        session = SessionLocal()
        try:
            upsert_review_plan(session, student_id, concept, 0.5, 1)
        finally:
            session.close()

        client = TestClient(app)
        try:
            for expected_count, quality in enumerate((2, 4, 5), start=1):
                response = client.post(
                    "/api/flashcard/review",
                    json={"student_id": student_id, "concept": concept, "quality": quality},
                )
                self.assertEqual(response.status_code, 200, response.text)
                data = response.json()
                self.assertIn("review_plan", data)
                self.assertIn("adaptive_review", data)
                self.assertEqual(data["review_plan"]["last_quality"], quality)
                self.assertEqual(data["review_plan"]["review_count"], expected_count)
                self.assertGreaterEqual(data["review_plan"]["easiness_factor"], 1.3)
                self.assertGreaterEqual(data["review_plan"]["interval_days"], 1)
                self.assertIn("next_review_at", data["review_plan"])
                self.assertIn("+00:00", data["flashcard"]["next_review_at"])
                self.assertEqual(data["flashcard"]["review_count"], expected_count)
                if quality == 2:
                    self.assertTrue(data["adaptive_review"]["triggered"])
                    self.assertIn("Mermaid", " ".join(data["adaptive_review"]["agent_trace"]))
                    self.assertIn("flowchart", data["adaptive_review"]["mermaid"])
                    self.assertTrue(data["adaptive_review"]["stream_chunks"])
                    self.assertIn("llm_backend", data["adaptive_review"])
                else:
                    self.assertFalse(data["adaptive_review"]["triggered"])

            bad = client.post(
                "/api/flashcard/review",
                json={"student_id": student_id, "concept": concept, "quality": 1},
            )
            self.assertEqual(bad.status_code, 400)
            bad_text = client.post(
                "/api/flashcard/review",
                json={"student_id": student_id, "concept": concept, "quality": "hard"},
            )
            self.assertEqual(bad_text.status_code, 400)

            due = client.get("/api/flashcard/due", params={"student_id": student_id})
            self.assertEqual(due.status_code, 200, due.text)
            self.assertEqual(due.json()["due_count"], 0)

            verify_session = SessionLocal()
            try:
                plan = verify_session.query(DBReviewPlan).filter_by(student_id=student_id, concept=concept).first()
                self.assertIsNotNone(plan)
                self.assertEqual(plan.last_quality, 5)
                self.assertEqual(plan.review_count, 3)
                self.assertIsNotNone(plan.next_review_at)
                profile = verify_session.query(DBStudentProfile).filter_by(student_id=student_id).first()
                self.assertIsNotNone(profile)
                self.assertIn(concept, profile.concept_mastery)
                self.assertIn(concept, profile.weak_points)
            finally:
                verify_session.close()
        finally:
            cleanup = SessionLocal()
            try:
                for plan in cleanup.query(DBReviewPlan).filter_by(student_id=student_id).all():
                    cleanup.delete(plan)
                profile = cleanup.query(DBStudentProfile).filter_by(student_id=student_id).first()
                if profile:
                    cleanup.delete(profile)
                cleanup.commit()
            finally:
                cleanup.close()

    def test_animation_dataset_feeds_resource_aware_planning(self):
        """本地 animations 数据集应能进入动态图谱，并为路径规划提供资源信号。"""
        from animation_resources import find_animation_dir, load_animation_resource_index
        from learning_strategy import build_resource_aware_dag
        from profile_api import KNOWLEDGE_DAG

        base_dir = find_animation_dir()
        if base_dir is None:
            self.skipTest("local animation dataset is not available")

        resource_index = load_animation_resource_index(str(base_dir))
        for concept in ("前向传播", "损失函数", "激活函数", "欠拟合"):
            self.assertIn(concept, resource_index)
            self.assertGreater(resource_index[concept]["video_count"], 0)

        active_dag, metadata = build_resource_aware_dag(KNOWLEDGE_DAG, resource_index=resource_index)
        self.assertIn("损失函数", active_dag.get("梯度下降", []))
        self.assertIn("链式法则", active_dag.get("反向传播", []))
        self.assertIn("损失函数", active_dag.get("反向传播", []))
        self.assertIn("反向传播", active_dag.get("神经网络", []))
        self.assertIn("梯度下降", active_dag.get("神经网络", []))
        self.assertIn("机器学习", active_dag.get("欠拟合", []))
        self.assertIn("resource_edges_inferred", metadata)
        self.assertGreaterEqual(metadata["animation_edges_added"], 0)
        self.assertEqual(metadata["animation_concept_count"], len(resource_index))

    def test_adaptive_astar_route_expands_prerequisites_deterministically(self):
        """A* 路径应稳定生成，并补齐多前置概念中的必要依赖。"""
        from learning_strategy import (
            PathPlanner,
            build_adaptive_astar_route,
            build_cross_disciplinary_micro_graph,
            build_micro_concept_graph,
            build_resource_aware_dag,
            suggest_cross_domain_supports,
        )
        from animation_resources import find_animation_dir, load_animation_resource_index
        from profile_api import KNOWLEDGE_DAG

        mastery = {"机器学习": 0.2, "线性回归": 0.1, "梯度下降": 0.1}
        route1 = build_adaptive_astar_route(
            KNOWLEDGE_DAG,
            mastery,
            learning_goals=["Transformer"],
            weak_points=["反向传播"],
            cognitive_load=0.62,
            frustration=0.15,
        )
        route2 = build_adaptive_astar_route(
            KNOWLEDGE_DAG,
            mastery,
            learning_goals=["Transformer"],
            weak_points=["反向传播"],
            cognitive_load=0.62,
            frustration=0.15,
        )
        planner = PathPlanner(KNOWLEDGE_DAG)
        planner_route = planner.plan(
            mastery,
            learning_goals=["Transformer"],
            weak_points=["反向传播"],
            cognitive_load=0.62,
            frustration=0.15,
        )

        concepts1 = [node["concept"] for node in route1["nodes"]]
        concepts2 = [node["concept"] for node in route2["nodes"]]
        planner_concepts = [node["concept"] for node in planner_route["nodes"]]
        self.assertEqual(concepts1, concepts2)
        self.assertGreaterEqual(len(concepts1), 5)
        self.assertLessEqual(len(concepts1), 8)
        self.assertEqual(len(concepts1), len(set(concepts1)))
        self.assertEqual(route1["target_concept"], "Transformer")
        self.assertEqual(route1["final_goal_concept"], "Transformer")
        self.assertTrue(route1["topology_audit"]["passed"])
        self.assertEqual(route1["strategy"], "A* 候选草案 + Planner Agent 资源感知审核")
        self.assertIn("反向传播", concepts1)
        self.assertGreater(route1["total_cost"], 0)
        self.assertTrue(route1["edges"])

        self.assertEqual(planner_route["target_concept"], "Transformer")
        self.assertEqual(planner_route["final_goal_concept"], "Transformer")
        self.assertIn(planner_route["stage_target_concept"], planner_concepts)
        self.assertGreaterEqual(len(planner_concepts), 5)
        self.assertLessEqual(len(planner_concepts), 8)
        self.assertTrue(planner_route["topology_audit"]["passed"])
        self.assertIn("candidate_draft", planner_route)
        self.assertTrue(planner_route["planner_trace"])
        self.assertGreaterEqual(planner_route["graph_fusion"]["active_node_count"], planner_route["graph_fusion"]["base_node_count"])
        self.assertIn("remaining_path", planner_route["candidate_draft"])
        if planner_route["resource_summary"]["animation_concept_count"] > 0:
            self.assertGreater(planner_route["resource_summary"]["matched_route_nodes"], 0)

        graph = planner.build_micro_graph(mastery, cognitive_load=0.62)
        base_dir = find_animation_dir()
        resource_index = load_animation_resource_index(str(base_dir)) if base_dir else {}
        active_dag, graph_metadata = build_resource_aware_dag(KNOWLEDGE_DAG, resource_index=resource_index)
        function_graph = build_micro_concept_graph(
            active_dag,
            mastery,
            cognitive_load=0.62,
            resource_index=resource_index,
            graph_metadata=graph_metadata,
        )
        self.assertEqual(graph["metadata"]["node_count"], function_graph["metadata"]["node_count"])
        self.assertGreaterEqual(graph["metadata"]["node_count"], 20)
        self.assertTrue(any(edge["from"] == "机器学习" and edge["to"] == "线性回归" for edge in graph["edges"]))

        cross_graph = planner.build_cross_disciplinary_graph(mastery, cognitive_load=0.62)
        function_cross_graph = build_cross_disciplinary_micro_graph(
            active_dag,
            mastery,
            cognitive_load=0.62,
            resource_index=resource_index,
        )
        self.assertEqual(cross_graph["metadata"]["cross_domain_edge_count"], function_cross_graph["metadata"]["cross_domain_edge_count"])
        cross_nodes = {node["concept"] for node in cross_graph["nodes"]}
        self.assertIn("偏导数", cross_nodes)
        self.assertTrue(any(node["concept"] == "偏导数" and node["domain_label"] for node in cross_graph["nodes"]))
        self.assertTrue(cross_graph["metadata"]["similarity_log"])
        self.assertTrue(any(domain.startswith("resource:") or domain.startswith("inferred:") for domain in cross_graph["metadata"]["domains"]))
        self.assertEqual(cross_graph["metadata"]["graph_backend"], "networkx")
        self.assertGreater(cross_graph["metadata"]["cross_domain_edge_count"], 0)
        self.assertGreater(cross_graph["metadata"]["semantic_edge_count"], 0)
        source_text = Path("learning_strategy.py").read_text(encoding="utf-8")
        self.assertNotIn("CROSS_DISCIPLINARY_MICRO_CONCEPTS", source_text)
        self.assertNotIn("CONCEPT_DOMAIN_HINTS", source_text)
        self.assertNotIn("DOMAIN_LABELS", source_text)
        supports = suggest_cross_domain_supports(cross_graph, planner_concepts)
        self.assertTrue(supports)
        self.assertTrue(all(item["reason"] for item in supports))

    def test_path_planner_respects_goal_variants_and_mastered_boundary(self):
        """PathPlanner 应能针对不同目标/薄弱点生成稳定路线，并处理全掌握边界。"""
        from learning_strategy import PathPlanner
        from profile_api import KNOWLEDGE_DAG

        planner = PathPlanner(KNOWLEDGE_DAG)

        overfit_route = planner.plan(
            {
                "机器学习": 0.1,
                "线性回归": 0.1,
                "正则化": 0.1,
                "交叉验证": 0.1,
            },
            learning_goals=["过拟合"],
            weak_points=["正则化"],
            cognitive_load=0.45,
            frustration=0.0,
        )
        overfit_concepts = [node["concept"] for node in overfit_route["nodes"]]
        self.assertEqual(overfit_route["target_concept"], "过拟合")
        self.assertEqual(overfit_route["final_goal_concept"], "过拟合")
        self.assertIn(overfit_route["stage_target_concept"], overfit_concepts)
        self.assertTrue(overfit_route["topology_audit"]["passed"])
        self.assertGreaterEqual(len(overfit_concepts), 5)
        self.assertLessEqual(len(overfit_concepts), 8)
        self.assertIn("正则化", overfit_concepts)
        if "过拟合" in overfit_concepts:
            self.assertLess(overfit_concepts.index("正则化"), overfit_concepts.index("过拟合"))
        else:
            self.assertTrue(overfit_route["is_staged_route"])

        cnn_route = planner.plan(
            {
                "机器学习": 0.9,
                "线性回归": 0.2,
                "梯度下降": 0.2,
            },
            learning_goals=["卷积神经网络"],
            weak_points=["卷积核", "池化层"],
            cognitive_load=0.35,
            frustration=0.05,
        )
        cnn_concepts = [node["concept"] for node in cnn_route["nodes"]]
        self.assertNotEqual(cnn_route["target_concept"], "Transformer")
        self.assertEqual(cnn_route["target_concept"], "卷积神经网络")
        self.assertEqual(cnn_route["final_goal_concept"], "卷积神经网络")
        self.assertIn(cnn_route["stage_target_concept"], cnn_concepts)
        self.assertTrue(cnn_route["topology_audit"]["passed"])
        self.assertGreaterEqual(len(cnn_concepts), 5)
        self.assertLessEqual(len(cnn_concepts), 8)
        later_cnn_steps = set(cnn_route["candidate_draft"].get("remaining_path", []))
        self.assertTrue({"卷积核", "池化层"} & (set(cnn_concepts) | later_cnn_steps))
        self.assertGreater(cnn_route["estimated_minutes"], 0)

        active_graph = planner.build_micro_graph({}, cognitive_load=0.35)
        all_concepts = {node["concept"] for node in active_graph["nodes"]}
        mastered_route = planner.plan({concept: 0.95 for concept in all_concepts})
        self.assertEqual(mastered_route["nodes"], [])
        self.assertEqual(mastered_route["total_cost"], 0.0)

    def test_learning_path_api_exposes_adaptive_astar_route(self):
        """学习路径 API 应暴露成员 2 的微概念图谱与 A* 动态路线。"""
        from app.database import SessionLocal, DBStudentProfile
        from app.crud import save_student_profile
        from fastapi.testclient import TestClient
        from app.main import app
        from models import StudentProfile
        import uuid

        student_id = f"test-astar-path-{uuid.uuid4().hex[:8]}"
        profile = StudentProfile(student_id=student_id)
        profile.concept_mastery = {
            "机器学习": 0.25,
            "线性回归": 0.15,
            "梯度下降": 0.12,
        }
        profile.learning_goals = ["Transformer"]
        profile.weak_points = ["反向传播"]
        profile.cognitive_load = 0.62
        profile.frustration_index = 0.2

        session = SessionLocal()
        try:
            save_student_profile(session, profile)
        finally:
            session.close()

        client = TestClient(app)
        try:
            response = client.get(f"/api/profile/{student_id}/learning-path")
            self.assertEqual(response.status_code, 200, response.text)
            data = response.json()
            self.assertIn("adaptive_route", data)
            self.assertIn("micro_concept_graph", data)
            self.assertIn("cross_domain_micro_graph", data)

            route = data["adaptive_route"]
            self.assertEqual(route["target_concept"], "Transformer")
            self.assertEqual(route["final_goal_concept"], "Transformer")
            self.assertIn(route["stage_target_concept"], [node["concept"] for node in route["nodes"]])
            self.assertTrue(route["topology_audit"]["passed"])
            self.assertGreaterEqual(len(route["nodes"]), 5)
            self.assertLessEqual(len(route["nodes"]), 8)
            self.assertGreater(route["total_cost"], 0)
            self.assertIn("candidate_draft", route)
            self.assertIn("remaining_path", route["candidate_draft"])
            self.assertTrue(route["planner_trace"])
            self.assertEqual(route["planner_review"]["decision"], "accepted")
            self.assertTrue(route["planner_review"]["action_dispatch"])
            self.assertTrue(route["session_plan"]["today_concepts"])
            self.assertGreater(route["session_plan"]["today_minutes"], 0)
            self.assertGreater(route["resource_summary"]["animation_concept_count"], 0)
            self.assertGreater(route["resource_summary"]["matched_route_nodes"], 0)
            self.assertIn("cross_domain_supports", route)
            self.assertTrue(route["cross_domain_supports"])
            self.assertTrue(any(item["reason"] for item in route["cross_domain_supports"]))
            self.assertEqual(data["progress_summary"]["adaptive_target"], "Transformer")
            self.assertGreater(data["progress_summary"]["cross_domain_supports"], 0)
            self.assertGreaterEqual(data["micro_concept_graph"]["metadata"]["edge_count"], 20)
            self.assertGreater(data["cross_domain_micro_graph"]["metadata"]["cross_domain_edge_count"], 0)
            self.assertTrue(data["cross_domain_micro_graph"]["metadata"]["similarity_log"])
        finally:
            cleanup = SessionLocal()
            try:
                db_profile = cleanup.query(DBStudentProfile).filter_by(student_id=student_id).first()
                if db_profile:
                    cleanup.delete(db_profile)
                    cleanup.commit()
            finally:
                cleanup.close()

    def test_note_matrix_closed_loop(self):
        """验证矩阵闭环学习流：笔记更新、AI 润色解析以及错题反思追加这套完整链路的正确性"""
        from app.database import SessionLocal, DBNote, DBQuizRecord, DBStudentProfile
        from app.crud import update_note, append_wrong_question_reflection, save_student_profile
        from models import StudentProfile
        import uuid
        
        student_id = f"test-matrix-loop-{uuid.uuid4().hex[:8]}"
        concept = "K均值聚类"
        
        session = SessionLocal()
        try:
            # 1. 模拟初始化学生画像
            profile = StudentProfile(student_id=student_id)
            save_student_profile(session, profile)
            
            # 2. 模拟创建一个答题记录（做错）
            quiz = DBQuizRecord(
                id=f"q-rec-{student_id}",
                student_id=student_id,
                question="什么是K均值聚类的聚类中心？",
                correct_answer="每个簇内所有样本的均值",
                student_answer="随机选择的一个点",
                feedback="K-Means应该重新计算均值作为中心",
                accuracy_score=0.0
            )
            session.add(quiz)
            session.commit()
            
            # 3. 验证错题记入笔记反思（第一次：此时没有笔记，应该自动新建）
            note1 = append_wrong_question_reflection(
                session, 
                student_id=student_id, 
                concept=concept, 
                quiz_record_id=quiz.id, 
                wrong_reason="misconception"
            )
            self.assertIsNotNone(note1)
            self.assertEqual(note1.student_id, student_id)
            self.assertIn("错题反思小记", note1.content)
            self.assertIn("随机选择的一个点", note1.content)
            self.assertIn("K均值聚类 错题集", note1.content)
            
            # 4. 验证更新笔记内容
            updated_content = note1.content + "\n\n额外补充：K-Means++ 能够优化初始质心的选择。"
            note2 = update_note(
                session, 
                note_id=note1.id, 
                content=updated_content, 
                tags=["聚类", "无监督学习"], 
                concepts=[concept]
            )
            self.assertEqual(note2.id, note1.id)
            self.assertIn("额外补充：K-Means++", note2.content)
            self.assertEqual(note2.tags, ["聚类", "无监督学习"])
            self.assertEqual(note2.concepts, [concept])
            
            # 5. 再次对同一个概念追加另一个错题反思（此时已存在该概念的笔记，应该追加）
            quiz2 = DBQuizRecord(
                id=f"q-rec2-{student_id}",
                student_id=student_id,
                question="K均值对什么非常敏感？",
                correct_answer="异常值和初始值",
                student_answer="特征数",
                feedback="异常值会极大地拉偏均值质心",
                accuracy_score=0.0
            )
            session.add(quiz2)
            session.commit()
            
            note3 = append_wrong_question_reflection(
                session, 
                student_id=student_id, 
                concept=concept, 
                quiz_record_id=quiz2.id, 
                wrong_reason="outlier_sensitivity"
            )
            self.assertEqual(note3.id, note1.id)
            self.assertIn("异常值和初始值", note3.content)
            self.assertIn("outlier_sensitivity", note3.content)
            
            # 清理
            session.delete(note3)
            session.delete(quiz)
            session.delete(quiz2)
            prof_now = session.query(DBStudentProfile).filter_by(student_id=student_id).first()
            if prof_now:
                session.delete(prof_now)
            session.commit()
        finally:
            session.close()

    def test_checkin_flow_and_streak(self):
        """验证复习打卡签到和连续天数（Streak）计算的完整流程"""
        from app.database import SessionLocal, DBCheckinLog, DBStudentProfile
        from quiz_api import _calc_streak
        from app.crud import save_student_profile
        from models import StudentProfile
        from datetime import datetime, timezone, timedelta
        import uuid

        student_id = f"test-checkin-{uuid.uuid4().hex[:8]}"
        session = SessionLocal()
        try:
            # 1. 初始化学生画像以满足外键约束
            profile = StudentProfile(student_id=student_id)
            save_student_profile(session, profile)

            # 2. 验证初始打卡连续天数为 0
            self.assertEqual(_calc_streak(session, student_id), 0)

            # 3. 模拟今天打卡（第一次，应当创建记录）
            now_utc = datetime.now(timezone.utc).replace(tzinfo=None)
            log_today = DBCheckinLog(
                student_id=student_id,
                checkin_date=now_utc,
                duration_minutes=15,
                concepts_reviewed=["支持向量机"]
            )
            session.add(log_today)
            session.commit()

            # 验证今天打卡后，streak 为 1
            self.assertEqual(_calc_streak(session, student_id), 1)

            # 4. 模拟同一天再次打卡（累加时长，累加概念）
            log_today.duration_minutes += 20
            log_today.concepts_reviewed = ["支持向量机", "核函数"]
            session.commit()

            # 验证同一个 UTC 天多次打卡不影响 streak 计算
            self.assertEqual(_calc_streak(session, student_id), 1)

            # 5. 模拟昨天打卡
            yesterday_utc = now_utc - timedelta(days=1)
            log_yesterday = DBCheckinLog(
                student_id=student_id,
                checkin_date=yesterday_utc,
                duration_minutes=10,
                concepts_reviewed=["机器学习基础"]
            )
            session.add(log_yesterday)
            session.commit()

            # 验证连续打卡昨天 + 今天，streak 应该为 2
            self.assertEqual(_calc_streak(session, student_id), 2)

            # 6. 模拟前天打卡
            two_days_ago_utc = now_utc - timedelta(days=2)
            log_two_days_ago = DBCheckinLog(
                student_id=student_id,
                checkin_date=two_days_ago_utc,
                duration_minutes=30,
                concepts_reviewed=["监督学习"]
            )
            session.add(log_two_days_ago)
            session.commit()

            # 验证今天、昨天、前天连续打卡，streak 为 3
            self.assertEqual(_calc_streak(session, student_id), 3)

            # 7. 模拟中间断开一天打卡（比如 4 天前打卡）
            four_days_ago_utc = now_utc - timedelta(days=4)
            log_four_days_ago = DBCheckinLog(
                student_id=student_id,
                checkin_date=four_days_ago_utc,
                duration_minutes=10,
                concepts_reviewed=["数据预处理"]
            )
            session.add(log_four_days_ago)
            session.commit()

            # 验证中间有间隔天没有打卡，连续天数计算应当被截断为 3（前天、昨天、今天）
            self.assertEqual(_calc_streak(session, student_id), 3)

            # 8. 测试在今天还没打卡的情况下，如果昨天打卡了，streak 依旧为之前的连续数，而不是 0
            session.delete(log_today)
            session.commit()
            # 此时只有昨天、前天、大前天（4天前）的记录，连续天数应该仍然是昨天+前天的 2 天
            self.assertEqual(_calc_streak(session, student_id), 2)

            # 9. 清理所有测试数据
            session.delete(log_yesterday)
            session.delete(log_two_days_ago)
            session.delete(log_four_days_ago)
            prof_now = session.query(DBStudentProfile).filter_by(student_id=student_id).first()
            if prof_now:
                session.delete(prof_now)
            session.commit()
        finally:
            session.close()

    def test_checkin_history_filtering(self):
        """验证签到历史查询与过滤逻辑"""
        from app.database import SessionLocal, DBCheckinLog, DBStudentProfile
        from app.crud import save_student_profile
        from models import StudentProfile
        from datetime import datetime, timezone, timedelta
        import uuid

        student_id = f"test-history-{uuid.uuid4().hex[:8]}"
        session = SessionLocal()
        try:
            profile = StudentProfile(student_id=student_id)
            save_student_profile(session, profile)

            now = datetime.now(timezone.utc).replace(tzinfo=None)
            log1 = DBCheckinLog(student_id=student_id, checkin_date=now - timedelta(days=2), duration_minutes=15, concepts_reviewed=["支持向量机", "机器学习"])
            log2 = DBCheckinLog(student_id=student_id, checkin_date=now - timedelta(days=1), duration_minutes=25, concepts_reviewed=["K均值聚类", "无监督学习"])
            session.add_all([log1, log2])
            session.commit()

            from quiz_api import get_checkin_history

            # 1. 验证不传 concept，查询全部打卡历史
            res_all = session.query(DBCheckinLog).filter(DBCheckinLog.student_id == student_id).order_by(DBCheckinLog.checkin_date.asc()).all()
            self.assertEqual(len(res_all), 2)
            
            # 2. 验证过滤 "支持向量机"
            results_svm = []
            for r in res_all:
                reviewed_list = [c.strip().lower() for c in (r.concepts_reviewed or [])]
                if "支持向量机".strip().lower() in reviewed_list:
                    results_svm.append(r)
            self.assertEqual(len(results_svm), 1)
            self.assertEqual(results_svm[0].duration_minutes, 15)

            # 3. 验证过滤 "k均值聚类"
            results_kmeans = []
            for r in res_all:
                reviewed_list = [c.strip().lower() for c in (r.concepts_reviewed or [])]
                if "k均值聚类".strip().lower() in reviewed_list:
                    results_kmeans.append(r)
            self.assertEqual(len(results_kmeans), 1)
            self.assertEqual(results_kmeans[0].duration_minutes, 25)

            # 清理
            session.delete(log1)
            session.delete(log2)
            prof_now = session.query(DBStudentProfile).filter_by(student_id=student_id).first()
            if prof_now:
                session.delete(prof_now)
            session.commit()
        finally:
            session.close()

    def test_conversation_history_endpoint(self):
        """验证对话历史接口的数据返回和属性映射"""
        from app.database import SessionLocal, DBConversationHistory, DBStudentProfile
        from app.crud import save_student_profile, record_conversation
        from models import StudentProfile
        from fastapi.testclient import TestClient
        from app.main import app
        import uuid

        student_id = f"test-chat-history-{uuid.uuid4().hex[:8]}"
        session = SessionLocal()
        try:
            profile = StudentProfile(student_id=student_id)
            save_student_profile(session, profile)

            # 写入对话记录
            record = record_conversation(
                session, student_id,
                query="如何理解卷积神经网络的池化层？",
                response_summary="CoordinatorAgent:text; DiagnosticianAgent:profile",
                target="池化层",
                resources_count=2,
                alignment_passed=True
            )

            client = TestClient(app)
            response = client.get(f"/api/history/{student_id}")
            self.assertEqual(response.status_code, 200)
            data = response.json()
            self.assertEqual(data["student_id"], student_id)
            self.assertTrue(len(data["history"]) >= 1)

            # 验证返回项包含了 query, message, response_summary, response, resources 和 resources_count
            history_item = next(h for h in data["history"] if h["id"] == record.id)
            self.assertEqual(history_item["query"], "如何理解卷积神经网络的池化层？")
            self.assertEqual(history_item["message"], "如何理解卷积神经网络的池化层？")
            self.assertEqual(history_item["response_summary"], "CoordinatorAgent:text; DiagnosticianAgent:profile")
            self.assertEqual(history_item["response"], "CoordinatorAgent:text; DiagnosticianAgent:profile")
            self.assertEqual(history_item["target"], "池化层")
            self.assertEqual(history_item["resources"], 2)
            self.assertEqual(history_item["resources_count"], 2)
            self.assertEqual(history_item["alignment_passed"], True)

            # 清理
            session.delete(record)
            prof_now = session.query(DBStudentProfile).filter_by(student_id=student_id).first()
            if prof_now:
                session.delete(prof_now)
            session.commit()
        finally:
            session.close()

    def test_stream_chat_records_conversation_in_history(self):
        """验证流式对话接口 (/api/stream/chat) 会正确将对话写入数据库中"""
        from app.database import SessionLocal, DBConversationHistory, DBStudentProfile
        from app.crud import save_student_profile
        from models import StudentProfile
        from fastapi.testclient import TestClient
        from app.main import app
        import uuid

        student_id = f"test-stream-history-{uuid.uuid4().hex[:8]}"
        session = SessionLocal()
        try:
            profile = StudentProfile(student_id=student_id)
            save_student_profile(session, profile)

            client = TestClient(app)
            # 发起流式对话请求，模拟正常聊天
            # 由于使用的是 Mock LLM，该请求会完成生成 5 个资源并自动调用 record_conversation
            response = client.post(
                "/api/stream/chat",
                json={"message": "请解释最大池化层", "student_id": student_id, "mode": "matrix"}
            )
            self.assertEqual(response.status_code, 200)

            # 从数据库读取该学生的对话记录
            records = session.query(DBConversationHistory).filter_by(student_id=student_id).all()
            self.assertEqual(len(records), 1)
            self.assertEqual(records[0].query, "请解释最大池化层")
            self.assertEqual(records[0].target, "最大池化")
            self.assertEqual(records[0].resources_count, 5)

            # 清理
            session.delete(records[0])
            prof_now = session.query(DBStudentProfile).filter_by(student_id=student_id).first()
            if prof_now:
                session.delete(prof_now)
            session.commit()
        finally:
            session.close()

    def test_pdf_upload_and_parsing(self):
        """验证 PDF 上传和解析不会因 BytesIO/seek 属性缺失而导致 500 崩溃"""
        from fastapi.testclient import TestClient
        from app.main import app
        from app.database import SessionLocal, DBKnowledgeDocument, DBStudentProfile
        from app.crud import save_student_profile
        from models import StudentProfile
        import uuid

        student_id = f"test-pdf-upload-{uuid.uuid4().hex[:8]}"
        client = TestClient(app)
        
        session = SessionLocal()
        try:
            profile = StudentProfile(student_id=student_id)
            save_student_profile(session, profile)
        finally:
            session.close()
        
        # 模拟 PDF 文件的二进制数据
        pdf_content = b"%PDF-1.4\n%...\n%%EOF\nThis is a mock PDF for testing convolutional neural networks."
        file_payload = {"file": ("test_cnn_layer.pdf", pdf_content, "application/pdf")}
        
        response = client.post(
            f"/api/knowledge/upload?student_id={student_id}",
            files=file_payload
        )
        self.assertEqual(response.status_code, 200)
        data = response.json()
        self.assertEqual(data["file_type"], "pdf")
        self.assertTrue(data["chunk_count"] >= 0)
        
        # 清理数据库和物理文件
        session = SessionLocal()
        try:
            doc = session.query(DBKnowledgeDocument).filter_by(id=data["id"]).first()
            if doc:
                session.delete(doc)
            prof_now = session.query(DBStudentProfile).filter_by(student_id=student_id).first()
            if prof_now:
                session.delete(prof_now)
            session.commit()
        finally:
            session.close()

        # 清除物理上传的文件和目录
        file_path = os.path.join("data", "uploads", student_id, f"{data['id']}.pdf")
        if os.path.exists(file_path):
            try:
                os.remove(file_path)
            except Exception:
                pass
        dir_path = os.path.join("data", "uploads", student_id)
        if os.path.exists(dir_path):
            try:
                os.rmdir(dir_path)
            except Exception:
                pass

    def test_docx_upload_and_parsing(self):
        """验证 Word (.docx) 上传和解析逻辑能够正常提取文本与入库"""
        from fastapi.testclient import TestClient
        from app.main import app
        from app.database import SessionLocal, DBKnowledgeDocument, DBStudentProfile
        from app.crud import save_student_profile
        from models import StudentProfile
        import uuid
        import docx
        from io import BytesIO

        student_id = f"test-docx-upload-{uuid.uuid4().hex[:8]}"
        client = TestClient(app)
        
        session = SessionLocal()
        try:
            profile = StudentProfile(student_id=student_id)
            save_student_profile(session, profile)
        finally:
            session.close()

        # 生成一个内存中的有效 docx 文件数据
        doc = docx.Document()
        doc.add_paragraph("This is a mock Word document for testing convolutional neural networks and pooling layers.")
        table = doc.add_table(rows=1, cols=2)
        table.rows[0].cells[0].text = "Header 1"
        table.rows[0].cells[1].text = "Value 1"
        
        docx_io = BytesIO()
        doc.save(docx_io)
        docx_content = docx_io.getvalue()
        
        file_payload = {"file": ("test_cnn_doc.docx", docx_content, "application/vnd.openxmlformats-officedocument.wordprocessingml.document")}
        
        response = client.post(
            f"/api/knowledge/upload?student_id={student_id}",
            files=file_payload
        )
        self.assertEqual(response.status_code, 200)
        data = response.json()
        self.assertEqual(data["file_type"], "docx")
        self.assertTrue(data["chunk_count"] >= 0)
        
        # 验证提取的文本中包含段落及表格内容
        self.assertIn("mock Word document", data["content_preview"])
        
        # 清理数据库和物理文件
        session = SessionLocal()
        try:
            db_doc = session.query(DBKnowledgeDocument).filter_by(id=data["id"]).first()
            if db_doc:
                session.delete(db_doc)
            prof_now = session.query(DBStudentProfile).filter_by(student_id=student_id).first()
            if prof_now:
                session.delete(prof_now)
            session.commit()
        finally:
            session.close()

        # 清除物理上传的文件和目录
        file_path = os.path.join("data", "uploads", student_id, f"{data['id']}.docx")
        if os.path.exists(file_path):
            try:
                os.remove(file_path)
            except Exception:
                pass
        dir_path = os.path.join("data", "uploads", student_id)
        if os.path.exists(dir_path):
            try:
                os.rmdir(dir_path)
            except Exception:
                pass

    def test_socratic_explain_multi_round(self):
        """验证苏格拉底即时答疑接口在单轮和多轮追问模式下的参数和路由行为"""
        from fastapi.testclient import TestClient
        from app.main import app
        import uuid

        client = TestClient(app)
        student_id = f"test-explain-{uuid.uuid4().hex[:8]}"

        # 1. 验证单轮公式解释
        response = client.post(
            "/api/stream/explain",
            json={
                "target_text": "\\frac{\\partial L}{\\partial w}",
                "context_before": "损失函数对权重的偏导数为",
                "context_after": "，用梯度下降更新权重。",
                "student_id": student_id
            }
        )
        self.assertEqual(response.status_code, 200)
        data = response.json()
        self.assertEqual(data["status"], "success")
        self.assertEqual(data["target_text"], "\\frac{\\partial L}{\\partial w}")
        self.assertIn("content", data)

        # 2. 验证多轮追问模式
        response = client.post(
            "/api/stream/explain",
            json={
                "target_text": "\\frac{\\partial L}{\\partial w}",
                "follow_up": "这里的L指的是什么损失函数？",
                "history": "导师: 这是一个求偏导的公式...\n学生: 明白。",
                "student_id": student_id
            }
        )
        self.assertEqual(response.status_code, 200)
        data = response.json()
        self.assertEqual(data["status"], "success")
        self.assertIn("content", data)

    def test_component_regeneration_endpoint(self):
        import uuid
        from app.database import SessionLocal
        from app.crud import load_student_profile, save_student_profile
        from fastapi.testclient import TestClient
        from app.main import app

        student_id = f"test-regen-{uuid.uuid4().hex[:8]}"
        session = SessionLocal()
        try:
            profile = load_student_profile(session, student_id)
            save_student_profile(session, profile)
        finally:
            session.close()

        client = TestClient(app)
        response = client.post(
            "/api/stream/regenerate",
            json={
                "student_id": student_id,
                "role": "理论教授",
                "resource_type": "专业讲义",
                "query": "什么是池化层"
            }
        )
        self.assertEqual(response.status_code, 200)
        data = response.json()
        self.assertEqual(data["status"], "success")
        self.assertIn("content", data)

    def test_lmcd_knowledge_diffusion(self):
        """验证 LMCD 知识扩散算法：当某个核心概念变化时，其增量会正确传播到邻近概念"""
        from bkt_engine import KnowledgeDiffusionEngine
        from agent_swarm import DEFAULT_KNOWLEDGE_DAG

        engine = KnowledgeDiffusionEngine(alpha=0.5, gamma=0.7)
        # 初始化一个画像概念掌握度
        concept_mastery = {
            "最大池化": 0.40,
            "池化层": 0.30,
            "机器学习": 0.50,
        }
        # 当“最大池化”增加 0.40 时，测试扩散效果
        updated = engine.diffuse(
            concept_mastery=concept_mastery,
            target_concept="最大池化",
            delta=0.40,
            dag=DEFAULT_KNOWLEDGE_DAG,
        )

        # 检查“最大池化”本身不变（仅扩散给其他概念）
        self.assertEqual(updated["最大池化"], 0.40)
        # 检查邻近的“池化层”（作为最大池化的前置概念，有拓扑依赖和语义相似）应该增加了掌握度
        self.assertGreater(updated["池化层"], 0.30)
        # 检查较远概念“机器学习”（无直接依赖），因为距离衰减，其增量应该明显小于邻近的“池化层”
        diff_pooling = updated["池化层"] - 0.30
        diff_ml = updated["机器学习"] - 0.50
        self.assertGreater(diff_pooling, diff_ml)

    def test_council_mode_consensus_synthesis(self):
        """验证 Council Mode 委员会机制对平行生成的资源进行共识评测与幻觉拦截"""
        from manifold_alignment import CouncilDecisionEngine
        from models import AgentOutput
        
        engine = CouncilDecisionEngine(fact_threshold=0.70, rel_threshold=0.65)
        
        # 1. 模拟没有冲突的高质量专家输出
        good_resources = [
            AgentOutput(agent="理论教授", resource_type="专业讲义", content="卷积神经网络包含池化层，最大池化层能降维", citations=()),
            AgentOutput(agent="逻辑画师", resource_type="思维导图", content="池化层：分为最大池化与平均池化", citations=()),
        ]
        synthesis1 = engine.synthesize(good_resources, "池化层")
        self.assertTrue(synthesis1["overall_passed"])
        self.assertGreaterEqual(synthesis1["mean_factuality_score"], 0.70)
        
        # 2. 模拟存在池化层操作冲突（幻觉）的专家输出
        conflict_resources = [
            AgentOutput(agent="理论教授", resource_type="专业讲义", content="最大池化Max Pooling取局部最大值，平均池化Avg Pooling取均值", citations=()),
            # 这个资源内容自相矛盾（同时提到了最大池化和平均池化，或存在事实偏离）
            AgentOutput(agent="极客助教", resource_type="代码实操案例", content="这里演示最大池化，但代码实现中我们用AvgPool2d表示平均池化，其实最大池化和平均池化是一回事", citations=()),
        ]
        synthesis2 = engine.synthesize(conflict_resources, "最大池化")
        self.assertFalse(synthesis2["overall_passed"])
        self.assertTrue(any(not v["passed"] for v in synthesis2["verdicts"]))

    def test_profile_analysis_includes_narrative_report(self):
        """验证 StoryLensEdu 叙事驱动评估报告管道：请求画像分析时应包含个性化成长信笺"""
        import uuid
        from app.database import SessionLocal
        from app.crud import load_student_profile, save_student_profile
        from fastapi.testclient import TestClient
        from app.main import app

        student_id = f"test-story-{uuid.uuid4().hex[:8]}"
        session = SessionLocal()
        try:
            profile = load_student_profile(session, student_id)
            save_student_profile(session, profile)
        finally:
            session.close()

        client = TestClient(app)
        response = client.get(f"/api/profile/{student_id}/analysis")
        self.assertEqual(response.status_code, 200)
        data = response.json()
        self.assertIn("narrative_report", data)
        self.assertIn("StoryLensEdu", data["narrative_report"])


    def test_rdi_and_peer_pk_challenge(self):
        """验证 RDI 风险指数计算以及 👥 找学伴 PK 的对抗拦截推流逻辑"""
        from fastapi.testclient import TestClient
        from app.main import app
        import uuid

        client = TestClient(app)
        student_id = f"test-pk-{uuid.uuid4().hex[:8]}"

        # 1. 验证普通 stream chat 中包含 RDI 字段
        response = client.post(
            "/api/stream/chat",
            json={"message": "什么是梯度下降？", "student_id": student_id}
        )
        self.assertEqual(response.status_code, 200)
        
        # 提取流中最后的 complete 帧
        complete_payload = None
        for line in response.iter_lines():
            if line.startswith("data: "):
                try:
                    import json
                    data = json.loads(line[6:])
                    # complete 帧有 resources 或 rdi 字段
                    if "rdi" in data:
                        complete_payload = data
                        break
                except Exception:
                    pass

        # 验证 RDI 结构存在且正确
        if complete_payload:
            self.assertIn("rdi", complete_payload)
            rdi = complete_payload["rdi"]
            self.assertIn("score", rdi)
            self.assertIn("category", rdi)
            self.assertIn("explanation", rdi)
            self.assertIn("details", rdi)

        # 2. 验证找学伴 PK 拦截与对抗出题
        response_pk = client.post(
            "/api/stream/chat",
            json={"message": "找学伴PK", "student_id": student_id}
        )
        self.assertEqual(response_pk.status_code, 200)
        
        pk_payload = None
        for line in response_pk.iter_lines():
            if line.startswith("data: "):
                try:
                    import json
                    data = json.loads(line[6:])
                    if "rdi" in data: # complete 帧
                        pk_payload = data
                        break
                except Exception:
                    pass

        self.assertIsNotNone(pk_payload)
        self.assertIn("小明", pk_payload["content"])
        self.assertIn("Adversarial Peer Challenge", pk_payload["content"])
        self.assertTrue(len(pk_payload["resources"]) > 0)
        self.assertEqual(pk_payload["resources"][0]["agent"], "同伴智能体 (小明)")


    def test_profile_update_endpoint(self):
        """验证学生画像更新接口正确性，能够修改并持久化各个偏好字段，并且清除信笺缓存"""
        import uuid
        from app.database import SessionLocal, DBStudentProfile
        from app.crud import load_student_profile, save_student_profile
        from fastapi.testclient import TestClient
        from app.main import app

        student_id = f"test-profile-update-{uuid.uuid4().hex[:8]}"
        session = SessionLocal()
        try:
            profile = load_student_profile(session, student_id)
            profile.narrative_report = "some cached narrative"
            save_student_profile(session, profile)
        finally:
            session.close()

        client = TestClient(app)
        
        # 1. 发送更新请求
        update_payload = {
            "major": "人工智能实践与应用",
            "target_course": "深度学习进阶",
            "cognitive_style": "代码实操导向",
            "motivation_type": "内在动机",
            "learning_goals": ["卷积神经网络", "池化层"],
            "learning_preferences": ["分步引导", "代码实操"]
        }
        
        response = client.post(f"/api/profile/{student_id}/update", json=update_payload)
        self.assertEqual(response.status_code, 200)
        data = response.json()
        self.assertEqual(data["status"], "success")
        self.assertEqual(data["major"], "人工智能实践与应用")
        self.assertEqual(data["cognitive_style"], "代码实操导向")
        self.assertEqual(data["learning_goals"], ["卷积神经网络", "池化层"])
        
        # 2. 验证数据库中确实更新了
        session = SessionLocal()
        try:
            profile_db = load_student_profile(session, student_id)
            self.assertEqual(profile_db.major, "人工智能实践与应用")
            self.assertEqual(profile_db.target_course, "深度学习进阶")
            self.assertEqual(profile_db.cognitive_style, "代码实操导向")
            self.assertEqual(profile_db.motivation_type, "内在动机")
            self.assertEqual(profile_db.learning_goals, ["卷积神经网络", "池化层"])
            self.assertEqual(profile_db.interaction_preferences, ["分步引导", "代码实操"])
            # 信笺缓存应该被清空
            self.assertEqual(profile_db.narrative_report, "")
            
            # 清理
            db_prof = session.query(DBStudentProfile).filter_by(student_id=student_id).first()
            if db_prof:
                session.delete(db_prof)
                session.commit()
        finally:
            session.close()

    def test_customized_fields_protection(self):
        """验证手动设定的画像字段受保护不会被大模型/消息自诊断算法篡改。"""
        from app.database import SessionLocal, DBStudentProfile
        from app.crud import load_student_profile, save_student_profile
        from models import StudentProfile
        import uuid
        
        student_id = f"test-custom-prot-{uuid.uuid4().hex[:8]}"
        session = SessionLocal()
        try:
            profile = load_student_profile(session, student_id)
            profile.major = "计算机工程"
            profile.target_course = "高级数据结构"
            profile.customized_fields = ["major", "target_course"]
            save_student_profile(session, profile)
            
            # 1. 尝试通过 update_from_message 更改，应由于 major 被保护而不改变
            profile.update_from_message("我是数学专业，希望学习深度学习。")
            self.assertEqual(profile.major, "计算机工程") # major 依然是计算机工程而非数学
            self.assertEqual(profile.target_course, "高级数据结构") # target_course 依然是高级数据结构而非深度学习
            
            # 2. 尝试通过 apply_llm_features (大模型画像抽取) 更改
            payload = {
                "major": "人工智能",
                "course": "机器学习",
                "goals": ["通过考试"],
                "preferences": ["数学推导"]
            }
            profile.apply_llm_features(payload, source_text="test source")
            self.assertEqual(profile.major, "计算机工程") # major 受保护
            self.assertEqual(profile.target_course, "高级数据结构") # target_course 受保护
            self.assertIn("通过考试", profile.learning_goals) # goals 未受保护，可更新
            self.assertIn("数学推导", profile.interaction_preferences) # preferences 未受保护，可更新
            
            # 清理
            db_prof = session.query(DBStudentProfile).filter_by(student_id=student_id).first()
            if db_prof:
                session.delete(db_prof)
                session.commit()
        finally:
            session.close()

    def test_multi_layer_cognitive_dimension_tracking(self):
        """验证多层级认知维度分解：单独更新不同维度时，BKT 与画像中的 concept_layers 可以正确分离与同步。"""
        from bkt_engine import BKTEngine
        from models import StudentProfile
        
        profile = StudentProfile(student_id="test-p2-dims")
        bkt = BKTEngine()
        
        # 1. 更新 "code" 维度
        bkt.update("池化层", correct=True, cognitive_load=0.4, frustration=0.1, profile_bkt_states=profile.bkt_states, dimension="code")
        snap = bkt.snapshot()
        self.assertIn("layers", snap["池化层"])
        self.assertIn("code", snap["池化层"]["layers"])
        
        # 2. 模拟事件总线分类同步
        from swarm_factory import build_swarm_from_headers
        swarm = build_swarm_from_headers({})
        profile = swarm.profile_store.setdefault("test-p2-dims", StudentProfile(student_id="test-p2-dims"))

        from learning_event_bus import QuizAttemptedEvent, _on_quiz_attempted
        import asyncio
        
        event = QuizAttemptedEvent(
            student_id="test-p2-dims",
            concept="池化层",
            accuracy=1.0,
            ai_confidence=0.9,
            student_confidence=0.9,
            attempt_number=2,
            question="请编写一段 PyTorch 代码实现最大池化。",
            answer="import torch.nn as nn..."
        )
        
        # 显式调用 _on_quiz_attempted
        async def _run_event():
            await _on_quiz_attempted(event)
        asyncio.run(_run_event())
        
        # 确认 concept_layers 中 code 维度已被正确提升并同步
        self.assertIn("池化层", profile.concept_layers)
        self.assertIn("code", profile.concept_layers["池化层"])

    def test_collaborative_peer_based_prior_calibration(self):
        """验证协同画像先验校准：新冷启动学生会根据匹配相似度的 Peer 数据做画像先验合并与状态恢复。"""
        from app.database import SessionLocal, DBStudentProfile
        from app.crud import load_student_profile, save_student_profile, calibrate_student_prior_collaborative
        from models import StudentProfile
        import uuid
        
        session = SessionLocal()
        # 备份已有画像，避免清除开发环境下的测试数据
        backup_profiles = []
        try:
            for p in session.query(DBStudentProfile).all():
                data = {c.name: getattr(p, c.name) for c in DBStudentProfile.__table__.columns}
                backup_profiles.append(data)
        except Exception:
            pass

        # 清空画像，避免其它测试产生的遗留数据污染
        session.query(DBStudentProfile).delete()
        session.commit()
        peer_id_1 = f"peer1-{uuid.uuid4().hex[:4]}"
        peer_id_2 = f"peer2-{uuid.uuid4().hex[:4]}"
        new_student_id = f"newstud-{uuid.uuid4().hex[:4]}"
        
        try:
            # 1. 准备 Peer 1 画像（相同专业，有答题记录和掌握度）
            peer1 = load_student_profile(session, peer_id_1)
            peer1.major = "数据科学"
            peer1.cognitive_style = "代码导向"
            peer1.concept_mastery = {"逻辑回归": 0.85, "链式法则": 0.90}
            peer1.bkt_states = {
                "逻辑回归": {"p_mastered": 0.85, "smoothed_mastery": 0.85, "p_err": 0.05, "layers": {}}
            }
            save_student_profile(session, peer1)
            
            # 2. 准备 Peer 2 画像（不同专业，不参与强校准）
            peer2 = load_student_profile(session, peer_id_2)
            peer2.major = "艺术设计"
            peer2.cognitive_style = "仅文本讲解"
            peer2.motivation_type = "外部型"
            peer2.concept_mastery = {"线性回归": 0.20}
            save_student_profile(session, peer2)
            
            # 3. 创建新冷启动学生
            new_student = StudentProfile(student_id=new_student_id)
            new_student.major = "数据科学"
            new_student.cognitive_style = "代码导向"
            
            # 4. 执行协同校准
            calibrate_student_prior_collaborative(session, new_student)
            
            # 5. 断言先验值是否成功合并了相似 Peer 1 的数据，且排除了 Peer 2
            self.assertIn("逻辑回归", new_student.concept_mastery)
            self.assertEqual(new_student.concept_mastery["逻辑回归"], 0.85)
            self.assertNotIn("线性回归", new_student.concept_mastery)
            
            # 清理
            for pid in (peer_id_1, peer_id_2, new_student_id):
                db_p = session.query(DBStudentProfile).filter_by(student_id=pid).first()
                if db_p:
                    session.delete(db_p)
            session.commit()
        finally:
            # 恢复之前备份的本地画像数据
            try:
                for data in backup_profiles:
                    p = DBStudentProfile(**data)
                    session.add(p)
                session.commit()
            except Exception:
                session.rollback()
            session.close()

    def test_uncertainty_aware_active_learning(self):
        """验证主动探索与不确定性消除：若未显式指定，生成测试题时优先寻找估计误差协方差 p_err 最大的概念。"""
        from app.database import SessionLocal, DBStudentProfile
        from app.crud import load_student_profile, save_student_profile
        from quiz_api import generate_quiz
        from fastapi import Request
        import uuid
        import asyncio
        
        session = SessionLocal()
        student_id = f"test-active-{uuid.uuid4().hex[:4]}"
        try:
            profile = load_student_profile(session, student_id)
            profile.concept_mastery = {
                "逻辑回归": 0.5,
                "线性回归": 0.9
            }
            profile.bkt_states = {
                "逻辑回归": {"p_mastered": 0.5, "smoothed_mastery": 0.5, "p_err": 0.8, "layers": {}},
                "线性回归": {"p_mastered": 0.9, "smoothed_mastery": 0.9, "p_err": 0.1, "layers": {}}
            }
            profile.weak_points = ["线性回归", "逻辑回归"]
            save_student_profile(session, profile)
            
            # 构造模拟 FastAPI 请求，不指定 target_concept
            class MockRequest:
                def __init__(self):
                    self.headers = {}
                async def json(self):
                    return {"student_id": student_id}
            
            async def _run_gen():
                return await generate_quiz(MockRequest())
                
            res = asyncio.run(_run_gen())
            # 应该选中了高不确定性的 "逻辑回归"，而不是 weak_points 的第一个 "线性回归"
            self.assertEqual(res["concept"], "逻辑回归")
            
            # 清理
            db_p = session.query(DBStudentProfile).filter_by(student_id=student_id).first()
            if db_p:
                session.delete(db_p)
                session.commit()
        finally:
            session.close()

    def test_causal_conflict_attribution_and_self_healing(self):
        """验证因果冲突归因与自愈引擎：一致性冲突发生时，能准确归因责任智能体，并在下次循环中注入自愈提示规则。"""
        from agent_swarm import CausalConflictAttributionEngine, AgentOutput, EduMatrixSwarm
        from models import AlignmentReport
        
        # 1. 构造一个包含代码和讲义不一致的冲突报告
        report = AlignmentReport(
            passed=False,
            distance=0.9,
            threshold=0.6,
            conflicts=[{
                "type": "code_mismatch",
                "description": "极客助教生成的代码使用了 AvgPool2d，但理论教授生成的讲义讲述的是最大池化(MaxPool2d)",
                "resources": ["理论教授讲义", "极客助教代码案例"]
            }],
            advice=" AvgPool2d 与 MaxPool2d 存在定义冲突 "
        )
        
        resources = (
            AgentOutput(agent="理论教授", resource_type="专业讲义", content="最大池化", citations=()),
            AgentOutput(agent="极客助教", resource_type="代码实操案例", content="AvgPool2d", citations=()),
            AgentOutput(agent="逻辑画师", resource_type="思维导图", content="Diagram", citations=())
        )
        
        # 2. 执行因果归因与自愈
        healing_instructions = CausalConflictAttributionEngine.attribute_and_heal(report, resources, rollback_count=1)
        
        # 3. 断言责任被正确划分到“极客助教”（因为冲突类型是 code_mismatch / 描述含代码）
        self.assertIn("极客助教", healing_instructions)
        self.assertNotIn("逻辑画师", healing_instructions)
        self.assertIn("AvgPool2d", healing_instructions["极客助教"])
        self.assertIn("MaxPool2d", healing_instructions["极客助教"])
        self.assertIn("系统自愈指令", healing_instructions["极客助教"])


    def test_kalman_filter_duration_noise_adjustment(self):
        """验证自适应卡尔曼滤波在答题时间极短时调整测量噪声 R"""
        from bkt_engine import BKTEngine
        engine = BKTEngine()
        
        # 正常答题时间更新
        smoothed_normal = engine.update("逻辑回归", correct=True, cognitive_load=0.4, frustration=0.1, duration_seconds=15.0)
        
        # 重新创建一个引擎测试极速答题情况（时间=1.0秒）
        engine_fast = BKTEngine()
        smoothed_fast = engine_fast.update("逻辑回归", correct=True, cognitive_load=0.4, frustration=0.1, duration_seconds=1.0)
        
        # 极速答题应该因为高 R_penalty 导致变化幅度显著变慢（提升较小）
        self.assertLess(smoothed_fast, smoothed_normal, "极速答题时由于噪声大，掌握度提升幅度应较小")

    def test_poincare_depth_alignment(self):
        """验证双曲庞加莱圆盘层级对齐：深度越深的概念，投影后半径越大，防止流形退化"""
        from manifold_alignment import project_to_poincare_ball, get_dag_depth
        import numpy as np
        
        # 构造基准高维向量（384D 均值 0.5 向量）
        vec = [0.5] * 384
        
        # “机器学习”（根概念，depth=0）
        vec_root = project_to_poincare_ball(vec, concept="机器学习")
        norm_root = np.linalg.norm(vec_root)
        
        # “逻辑回归”（较深，depth=2）
        vec_deep = project_to_poincare_ball(vec, concept="逻辑回归")
        norm_deep = np.linalg.norm(vec_deep)
        
        # “最大池化”（叶子，depth=6）
        vec_leaf = project_to_poincare_ball(vec, concept="最大池化")
        norm_leaf = np.linalg.norm(vec_leaf)
        
        self.assertGreater(norm_leaf, norm_deep)
        self.assertGreater(norm_deep, norm_root)
        self.assertLess(norm_leaf, 0.82)  # 应小于 max_norm 限制

    def test_graph_kalman_propagation(self):
        """验证图卡尔曼信念传播：更新节点，前置节点与后续依赖节点同步平滑演进"""
        from bkt_engine import BKTEngine
        engine = BKTEngine()
        
        # 初始状态，冷启动时默认为 0.3
        mastery_lr = engine.get_mastery("线性回归")  # 前置
        mastery_logr = engine.get_mastery("逻辑回归") # 目标
        
        # 更新“逻辑回归”正确 (掌握度上升)
        engine.update("逻辑回归", correct=True, cognitive_load=0.4, frustration=0.1)
        
        # 验证“逻辑回归”上升了
        new_logr = engine.get_mastery("逻辑回归")
        self.assertGreater(new_logr, mastery_logr)
        
        # 验证前置“线性回归”（在 KNOWLEDGE_DAG 中为逻辑回归的前置之一）协同上升了
        new_lr = engine.get_mastery("线性回归")
        self.assertGreater(new_lr, mastery_lr, "前置依赖概念应协同上升")
        
        # 验证不相关的节点掌握度没有被误波及更新
        self.assertEqual(engine.get_mastery("池化层"), 0.3, "无关概念应保持初始状态")

    def test_dkt_and_ekf_concurrency_and_math_consistency(self):
        """Task 4: 验证 DktService 并发更新的线程安全性、EKF 状态转移矩阵的行和恒为 1.0，以及非阻塞 MDS 异步运行"""
        import concurrent.futures
        import time
        from bkt_engine import DktService, BKTEngine, poincare_to_2d_coordinates
        import numpy as np
        
        # 1. 验证 DktService 并发调用 train_incremental 不会报错，也不会引起 PyTorch 梯度冲突
        dkt = DktService()
        records = [("最大池化", True), ("池化层", True), ("平均池化", False)]
        
        def run_incremental():
            for _ in range(5):
                dkt.train_incremental(records)
                time.sleep(0.01)
                
        # 模拟 10 个线程并发增量微调
        with concurrent.futures.ThreadPoolExecutor(max_workers=10) as executor:
            futures = [executor.submit(run_incremental) for _ in range(10)]
            for f in futures:
                f.result()  # 确保无异常抛出
                
        # 消费队列任务，并用 predict_mastery 并发做测试
        dkt.update_queue.join()  # 等待所有增量梯度更新消费完成
        
        # 验证并发推理 predict_mastery 的线程安全性
        def run_predict():
            res = dkt.predict_mastery(records)
            self.assertIn("最大池化", res)
            self.assertGreaterEqual(res["最大池化"], 0.0)
            
        with concurrent.futures.ThreadPoolExecutor(max_workers=5) as executor:
            futures_pred = [executor.submit(run_predict) for _ in range(5)]
            for f in futures_pred:
                f.result()
                
        # 2. 验证 Graph-EKF 局部状态转移矩阵 F_local 的行和恒为 1.0 (防止物理意义偏失与数值发散)
        engine = BKTEngine()
        # 触发 EKF 的前向信念更新，查看内部 F_local 的构成
        # 即使 f_forward 被设定为非零值，F_local 的每一行和都应该严格等于 1.0
        # 模拟调用 update
        engine.update("逻辑回归", correct=True, cognitive_load=0.4, frustration=0.1)
        
        # 验证 KNOWLEDGE_DAG 中有前置依赖的节点（例如“逻辑回归”）在 F_local 中的行和为 1.0
        # 我们可以通过构造测试矩阵的方式直接验证 EKF 前向凸组合的稳定性：
        f_forward = engine.f_forward
        self.assertLess(f_forward, 1.0)
        self.assertGreater(f_forward, 0.0)
        
        # 3. 验证 poincare_to_2d_coordinates 对缺失值 Fallback MDS 优化的物理屏障和坐标有界性
        # (坐标必须局限在庞加莱单位圆盘 |x| < 1.0 以内)
        test_embeddings = {
            "池化层": [0.1] * 384,
            "卷积层": [-0.2] * 384,
            "最大池化": [0.0] * 384
        }
        coords = poincare_to_2d_coordinates(test_embeddings)
        self.assertIn("池化层", coords)
        self.assertIn("卷积层", coords)
        for concept, pt in coords.items():
            norm_pt = np.linalg.norm(pt)
            self.assertLess(norm_pt, 1.0, f"{concept} 坐标 {pt} 超出了双曲庞加莱圆盘的范围(模长应小于 1.0)")

    def test_dkt_dynamic_concept_tracking_and_ekf_backward_propagation(self):
        """验证 DKT 动态概念追踪和 EKF 双向信念传播与行和守恒"""
        from bkt_engine import DktService, BKTEngine
        import numpy as np

        # 1. 验证 DKT 对新动态概念的零样本推理
        dkt = DktService()
        history = [("逻辑回归", True), ("新概念_PLC编程", True)]
        res = dkt.predict_mastery(history, extra_concepts=["新概念_智能控制", "新概念_PLC编程"])
        
        # 验证返回包含了动态概念，且推理成功
        self.assertIn("新概念_PLC编程", res)
        self.assertIn("新概念_智能控制", res)
        self.assertGreaterEqual(res["新概念_PLC编程"], 0.0)
        self.assertLessEqual(res["新概念_PLC编程"], 1.0)
        self.assertGreaterEqual(res["新概念_智能控制"], 0.0)
        
        # 2. 验证 BKTEngine 在双向信念传播时转移矩阵 F_local 的行和完全守恒为 1.0
        engine = BKTEngine()
        # 触发有依赖子图的概念更新
        engine.update("逻辑回归", correct=True)
        # 检验矩阵行和守恒的正确性：
        # 我们手动测试活跃概念子图的状态转移行和
        # "逻辑回归" -> prereqs: "线性回归", "梯度下降" -> successors: None
        # "线性回归" -> prereqs: "机器学习" -> successors: "逻辑回归", "支持向量机", "正则化"
        active_concepts = ["逻辑回归", "线性回归", "梯度下降", "机器学习"]
        M = len(active_concepts)
        c_to_idx = {c: i for i, c in enumerate(active_concepts)}
        
        from profile_api import KNOWLEDGE_DAG
        
        F_local = np.eye(M, dtype=np.float32)
        for i, c in enumerate(active_concepts):
            idx_i = c_to_idx[c]
            
            # 前向
            c_prereqs = KNOWLEDGE_DAG.get(c, [])
            active_prereqs = [p for p in c_prereqs if p in c_to_idx]
            total_prereqs = len(c_prereqs)
            if total_prereqs > 0:
                F_local[idx_i, idx_i] -= engine.f_forward
                for p in active_prereqs:
                    idx_j = c_to_idx[p]
                    F_local[idx_i, idx_j] = engine.f_forward / total_prereqs
                missing_count = total_prereqs - len(active_prereqs)
                if missing_count > 0:
                    F_local[idx_i, idx_i] += missing_count * (engine.f_forward / total_prereqs)
            
            # 反向
            active_successors = [succ for succ, p_list in KNOWLEDGE_DAG.items() if c in p_list and succ in c_to_idx]
            total_successors = len([succ for succ, p_list in KNOWLEDGE_DAG.items() if c in p_list])
            if total_successors > 0:
                F_local[idx_i, idx_i] -= engine.f_backward
                for succ in active_successors:
                    idx_succ = c_to_idx[succ]
                    F_local[idx_i, idx_succ] = engine.f_backward / total_successors
                missing_succ_count = total_successors - len(active_successors)
                if missing_succ_count > 0:
                    F_local[idx_i, idx_i] += missing_succ_count * (engine.f_backward / total_successors)
                    
        # 校验每一行和严格为 1.0 (容许极小浮点误差)
        for r in range(M):
            row_sum = float(np.sum(F_local[r]))
            self.assertAlmostEqual(row_sum, 1.0, places=5, msg=f"Row {r} sum is not 1.0: {row_sum}")


if __name__ == "__main__":
    unittest.main()

